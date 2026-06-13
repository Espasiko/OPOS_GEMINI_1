import os
import re
import sys
import json
import time
import logging
import subprocess
from datetime import datetime
from logging.handlers import TimedRotatingFileHandler
from pathlib import Path
from typing import Optional

import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from fastapi import FastAPI, Request
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from mistralai import Mistral

# Blindaje UTF-8 en Windows: aunque .bat haga chcp 65001, Python 3.13 sobre cmd.exe
# puede caer a cp1252 si no se reconfigura. Esto evita UnicodeEncodeError con emojis.
if sys.platform.startswith("win"):
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

# ==========================================
# 🔐 CARGA DE VARIABLES DE ENTORNO (.env)
# ==========================================
# Compatible con script normal (.py) y con bundle PyInstaller (.exe).
# Para PyInstaller --onefile, sys.executable apunta al .exe en disco;
# __file__ apuntaría al dir temporal _MEIxxx (no útil para el usuario).
if getattr(sys, "frozen", False):
    APP_DIR = Path(sys.executable).resolve().parent
else:
    APP_DIR = Path(__file__).resolve().parent

# Probamos varios nombres por compatibilidad: .env (paquete Nina) y .env.backend (repo dev)
for _candidate in (".env", ".env.backend"):
    _env_path = APP_DIR / _candidate
    if _env_path.exists():
        load_dotenv(_env_path)
        break

# ==========================================
# 📝 LOGGING: stdout + archivo persistente con rotación diaria
# ==========================================
# Crea logs/agente_YYYY-MM-DD.log en el mismo dir del .exe (o del .py).
# Rota cada medianoche, conserva 14 días para diagnóstico remoto.
_LOG_DIR = APP_DIR / "logs"
try:
    _LOG_DIR.mkdir(exist_ok=True)
except Exception:
    _LOG_DIR = APP_DIR  # fallback si no se puede crear (permisos)

_LOG_FORMAT = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")

_root_logger = logging.getLogger()
_root_logger.setLevel(logging.INFO)
# Limpia handlers previos por si hubiera (evita duplicados al recargar)
for _h in list(_root_logger.handlers):
    _root_logger.removeHandler(_h)

# Handler 1: consola (la ventana negra que ve el usuario)
_console_handler = logging.StreamHandler(sys.stdout)
_console_handler.setFormatter(_LOG_FORMAT)
_root_logger.addHandler(_console_handler)

# Handler 2: archivo rotatorio diario
try:
    _file_handler = TimedRotatingFileHandler(
        _LOG_DIR / "agente.log",
        when="midnight",
        interval=1,
        backupCount=14,           # conserva 14 días
        encoding="utf-8",
        utc=False,
    )
    _file_handler.suffix = "%Y-%m-%d"  # agente.log.2026-04-27
    _file_handler.setFormatter(_LOG_FORMAT)
    _root_logger.addHandler(_file_handler)
except Exception as _e:
    # Si falla (disco lleno, permisos), seguimos solo con consola
    print(f"[WARN] No se pudo abrir log file: {_e}")

log = logging.getLogger("agente-escritor")
log.info(f"📁 Logs persistentes en: {_LOG_DIR}")

MISTRAL_API_KEY = os.getenv("MISTRAL_API_KEY", "").strip()
OBSIDIAN_API_KEY = os.getenv("OBSIDIAN_REST_API_KEY", "").strip()
OBSIDIAN_REST_URL_ENV = os.getenv("OBSIDIAN_REST_URL", "").strip()
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY", "").strip()
MISTRAL_MODEL = os.getenv("MISTRAL_AGENT_MODEL", "mistral-large-latest")
PROXY_PORT = int(os.getenv("PROXY_PORT", "8000"))
# SEGURIDAD: por defecto escuchamos SOLO en localhost (127.0.0.1). Las tools pueden
# leer/escribir/borrar el vault, así que NO conviene exponer el proxy a la red. Si alguien
# necesita acceso desde otra máquina (raro), puede poner PROXY_HOST=0.0.0.0 en .env.
PROXY_HOST = os.getenv("PROXY_HOST", "127.0.0.1").strip() or "127.0.0.1"
OBSIDIAN_PORT = 27123
# Raíz del vault en disco. Por defecto = carpeta del .exe (que vive en la raíz del
# vault de Nina). Permite override con VAULT_PATH en .env. Usado por git_* y read_pdf.
VAULT_PATH = os.getenv("VAULT_PATH", "").strip() or str(APP_DIR)
# Ejecutable git. Por defecto el 'git' del sistema; si se entrega MinGit portable junto al
# .exe, poner GIT_EXE=mingit\cmd\git.exe en .env (se resuelve relativo a la carpeta del .exe).
GIT_EXE = os.getenv("GIT_EXE", "").strip() or "git"
if GIT_EXE != "git" and not os.path.isabs(GIT_EXE):
    _git_cand = APP_DIR / GIT_EXE
    if _git_cand.exists():
        GIT_EXE = str(_git_cand)

if not MISTRAL_API_KEY:
    log.warning("⚠️  MISTRAL_API_KEY no configurada en .env.backend. El proxy no podrá llamar a Mistral.")
if not OBSIDIAN_API_KEY:
    log.warning("⚠️  OBSIDIAN_REST_API_KEY no configurada. Las herramientas de vault fallarán.")

app = FastAPI(title="Agente EscritorAIA Proxy")


def detectar_ip_host_desde_wsl() -> Optional[str]:
    """Solo en WSL: detecta la IP del host Windows usando `ip route`.
    En Windows nativo no se llama (Obsidian REST está en 127.0.0.1)."""
    try:
        out = subprocess.check_output(["ip", "route"], text=True, timeout=2)
        for line in out.splitlines():
            if line.startswith("default via "):
                return line.split()[2]
    except Exception as e:
        log.warning(f"No pude detectar IP host via 'ip route': {e}")
    return None


def get_obsidian_base_url() -> str:
    """Resuelve la URL del REST API de Obsidian.
    Prioridad:
      1) OBSIDIAN_REST_URL en .env (si está configurada).
      2) En Windows nativo: 127.0.0.1 (mismo equipo que Obsidian).
      3) En Linux/WSL: detectar IP del host Windows con `ip route`.
      4) Fallback: localhost.
    """
    if OBSIDIAN_REST_URL_ENV:
        return OBSIDIAN_REST_URL_ENV
    # Windows nativo: el .exe corre en la misma máquina que Obsidian
    if sys.platform.startswith("win"):
        return f"http://127.0.0.1:{OBSIDIAN_PORT}"
    # Linux/WSL: hay que cruzar a Windows host
    ip = detectar_ip_host_desde_wsl()
    if ip:
        return f"http://{ip}:{OBSIDIAN_PORT}"
    return f"http://127.0.0.1:{OBSIDIAN_PORT}"


OBSIDIAN_URL = get_obsidian_base_url()
log.info(f"🔗 Obsidian URL: {OBSIDIAN_URL}")
log.info(f"🤖 Mistral model: {MISTRAL_MODEL}")

HEADERS = {
    "Authorization": f"Bearer {OBSIDIAN_API_KEY}",
    "Content-Type": "text/markdown",
}

# ==========================================
# 🛠️ LA MAGIA: DEFINICIÓN DE HERRAMIENTAS
# ==========================================

def read_obsidian_note(filename: str) -> str:
    """Lee el contenido de una nota de Obsidian dada su ruta relativa."""
    filename = filename.replace(" ", "%20")
    if not filename.endswith(".md"): filename += ".md"
    
    url = f"{OBSIDIAN_URL}/vault/{filename}"
    try:
        r = requests.get(url, headers={"Authorization": HEADERS["Authorization"]})
        if r.status_code == 200:
            return r.text
        return f"Error: No se encontró la nota (Código {r.status_code})"
    except Exception as e:
        return f"Error de conexión: {str(e)}"

def create_obsidian_note(filename: str, content: str) -> str:
    """Crea un fichero nuevo Markdown y escribe el contenido dentro del vault."""
    filename = filename.replace(" ", "%20")
    if not filename.endswith(".md"): filename += ".md"

    url = f"{OBSIDIAN_URL}/vault/{filename}"
    try:
        r = requests.put(url, headers=HEADERS, data=content.encode('utf-8'))
        if r.status_code in [200, 201, 204]:
            return f"¡Éxito! La nota {filename} fue creada perfectamente."
        log.warning(f"⚠️ create_obsidian_note({filename}) → HTTP {r.status_code}: {r.text[:200]}")
        return f"Fallo al crear nota. Código {r.status_code}: {r.text[:100]}"
    except Exception as e:
        return f"Error de conexión: {str(e)}"

def update_obsidian_note(filename: str, content: str) -> str:
    """Añade (append) contenido al final de una nota Markdown existente."""
    filename = filename.replace(" ", "%20")
    if not filename.endswith(".md"): filename += ".md"

    url = f"{OBSIDIAN_URL}/vault/{filename}"
    try:
        # POST para append
        r = requests.post(url, headers=HEADERS, data=content.encode('utf-8'))
        if r.status_code in [200, 201, 204]:
            return f"¡Éxito! Texto añadido a {filename}."
        log.warning(f"⚠️ update_obsidian_note({filename}) → HTTP {r.status_code}: {r.text[:200]}")
        return f"Fallo al editar nota. Código {r.status_code}: {r.text[:100]}"
    except Exception as e:
        return f"Error de conexión: {str(e)}"


def overwrite_obsidian_note(filename: str, content: str) -> str:
    """Reescribe completamente una nota existente con contenido nuevo.
    CUIDADO: esto borra todo el contenido anterior de la nota."""
    filename = filename.replace(" ", "%20")
    if not filename.endswith(".md"): filename += ".md"

    url = f"{OBSIDIAN_URL}/vault/{filename}"
    try:
        r = requests.put(url, headers=HEADERS, data=content.encode('utf-8'))
        if r.status_code in [200, 201, 204]:
            return f"¡Éxito! La nota {filename} fue reescrita completamente."
        log.warning(f"⚠️ overwrite_obsidian_note({filename}) → HTTP {r.status_code}: {r.text[:200]}")
        return f"Fallo al reescribir nota. Código {r.status_code}: {r.text[:100]}"
    except Exception as e:
        return f"Error de conexión: {str(e)}"


def delete_obsidian_note(filename: str) -> str:
    """Borra una nota del vault de Obsidian. La nota se mueve a la papelera de Obsidian."""
    filename = filename.replace(" ", "%20")
    if not filename.endswith(".md"): filename += ".md"

    url = f"{OBSIDIAN_URL}/vault/{filename}"
    try:
        r = requests.delete(url, headers={"Authorization": HEADERS["Authorization"]})
        if r.status_code in [200, 204]:
            return f"¡Éxito! La nota {filename} fue eliminada."
        elif r.status_code == 404:
            return f"Error: La nota {filename} no existe."
        return f"Fallo al borrar nota. Código {r.status_code}"
    except Exception as e:
        return f"Error de conexión: {str(e)}"


def list_vault_files(folder: str = "/", extension: str = "") -> str:
    """Lista los archivos del vault de Obsidian, opcionalmente filtrados por carpeta y extensión.

    BUGFIX 28/05/2026: la Obsidian Local REST API NO es recursiva en /vault/.
    Si el cliente pide una subcarpeta, hay que llamar a /vault/{folder}/ directamente.
    Si no se especifica carpeta, recorremos recursivamente para tener un listado útil.

    CAMBIO 01/06/2026: por defecto NO filtra por extensión → lista TODO (md, docx, pdf, txt, etc.)
    para que el agente se entere de archivos no-md que el usuario pueda haber pegado y los pase
    por `ingest_file()`. Si quieres solo md, pasa extension='.md'. Admite varias separadas por
    coma: extension='.md,.docx,.pdf'.

    OJO: La REST API de Obsidian respeta el filtro `showUnsupportedFiles` del vault (`app.json`).
    Si está en false, los .docx/.txt no aparecen aunque existan en disco. Para esos casos,
    usa `scan_disk_files(folder)` que lee el filesystem directamente.
    """
    headers = {"Authorization": HEADERS["Authorization"], "Accept": "application/json"}

    def _fetch(path: str):
        """Llama a /vault/{path}/ y devuelve la lista cruda de entradas (archivos y carpetas)."""
        path_clean = path.strip("/")
        url = f"{OBSIDIAN_URL}/vault/" if not path_clean else f"{OBSIDIAN_URL}/vault/{path_clean}/"
        r = requests.get(url, headers=headers, timeout=10)
        if r.status_code != 200:
            return None, f"HTTP {r.status_code} en {url}"
        return r.json().get("files", []), None

    def _walk(base: str, depth: int = 0, max_depth: int = 4) -> list:
        """Recorre subcarpetas hasta max_depth y devuelve rutas relativas a la raíz del vault."""
        out = []
        entries, err = _fetch(base)
        if err or entries is None:
            return out
        base_clean = base.strip("/")
        for entry in entries:
            full = f"{base_clean}/{entry}" if base_clean else entry
            if entry.endswith("/"):
                if depth < max_depth:
                    out.extend(_walk(full.rstrip("/"), depth + 1, max_depth))
            else:
                out.append(full)
        return out

    try:
        if folder and folder != "/":
            files = _walk(folder)
        else:
            files = _walk("")

        if extension:
            exts = tuple(e.strip().lower() for e in extension.split(",") if e.strip())
            files = [f for f in files if f.lower().endswith(exts)]

        if not files:
            tip = ""
            if not extension:
                tip = (" Si esperabas ver un archivo no-md (docx/pdf/txt), comprueba que "
                       "Obsidian tenga `showUnsupportedFiles: true` en app.json, o usa "
                       "`scan_disk_files()` que lee el disco directo.")
            return f"No se encontraron archivos en '{folder}'.{tip}"
        return json.dumps(files[:100], ensure_ascii=False, indent=2)
    except Exception as e:
        return f"Error de conexión: {str(e)}"


def scan_disk_files(folder: str = "", extension: str = "") -> str:
    """Lista archivos del vault leyendo DIRECTAMENTE del disco (sin pasar por REST API).

    Útil cuando Obsidian filtra .docx/.txt/.pdf con `showUnsupportedFiles: false`.
    Recorre recursivamente con un max_depth=6 y devuelve rutas relativas al vault.
    Ignora carpetas internas (`.obsidian`, `.git`, `.smart-env`, `node_modules`, etc.).

    Args:
        folder: subcarpeta del vault (vacío = raíz).
        extension: filtro por extensión (vacío = todas). Admite coma-separada: '.md,.docx'.
    """
    import os.path as _osp
    vroot = _osp.realpath(VAULT_PATH)
    base = _osp.realpath(_osp.join(vroot, folder.strip("/"))) if folder else vroot
    if not base.startswith(vroot):
        return "Error: por seguridad solo puedo escanear dentro del vault."
    if not _osp.isdir(base):
        return f"Error: la carpeta '{folder}' no existe en disco."

    exts = tuple(e.strip().lower() for e in extension.split(",") if e.strip()) if extension else None
    SKIP_DIRS = {".obsidian", ".git", ".smart-env", "node_modules", ".trash", "logs", "mingit"}
    out = []
    for root, dirs, files in os.walk(base):
        dirs[:] = [d for d in dirs if d not in SKIP_DIRS and not d.startswith(".")]
        depth = root[len(vroot):].count(os.sep)
        if depth > 6:
            dirs[:] = []
            continue
        for fname in files:
            if fname.startswith("."):
                continue
            if exts and not fname.lower().endswith(exts):
                continue
            rel = _osp.relpath(_osp.join(root, fname), vroot).replace("\\", "/")
            out.append(rel)
            if len(out) >= 300:
                break
        if len(out) >= 300:
            break

    if not out:
        return f"No se encontraron archivos en disco en '{folder}'."
    return json.dumps(sorted(out), ensure_ascii=False, indent=2)


def ingest_file(path: str, target_md: str = "") -> str:
    """Extrae el texto de un archivo no-md (.docx, .pdf, .txt, .csv, .rtf) del vault
    y crea una nota .md GEMELA al lado con el contenido extraído. Esto permite que
    Smart Connections / smart-lookup indexen semánticamente el contenido.

    Args:
        path: ruta del archivo dentro del vault (ej '03_NOTAS/borrador.docx').
        target_md: ruta destino del .md (opcional). Si está vacío, usa '{path}.md'.

    Smart Connections detectará el .md nuevo en cuestión de segundos y lo añadirá
    al embed queue automáticamente.
    """
    import os.path as _osp
    full = _osp.realpath(path if _osp.isabs(path) else _osp.join(VAULT_PATH, path))
    vroot = _osp.realpath(VAULT_PATH)
    if not (full == vroot or full.startswith(vroot + os.sep)):
        return "Error: por seguridad solo puedo ingerir archivos dentro del vault."
    if not _osp.exists(full):
        return f"Error: no se encontró el archivo '{path}' en disco."

    ext = _osp.splitext(full)[1].lower()
    if ext == ".md":
        return f"'{path}' ya es markdown — Smart Connections lo indexa automáticamente sin ingest_file."

    texto = ""
    if ext == ".docx":
        try:
            import docx
            d = docx.Document(full)
            partes = [p.text for p in d.paragraphs if p.text and p.text.strip()]
            texto = "\n".join(partes)
        except Exception as e:
            return f"Error leyendo .docx: {e}"
    elif ext == ".pdf":
        rel = _osp.relpath(full, vroot).replace("\\", "/")
        texto = read_pdf(rel, max_pages=100)
        if texto.startswith("PDF "):
            try:
                texto = texto.split("\n\n", 1)[1]
            except Exception:
                pass
        if texto.startswith("Error"):
            return texto
    elif ext in (".txt", ".csv", ".log", ".tsv", ".json", ".yaml", ".yml"):
        try:
            with open(full, "r", encoding="utf-8", errors="replace") as f:
                texto = f.read()
        except Exception as e:
            return f"Error leyendo {ext}: {e}"
    elif ext == ".rtf":
        try:
            with open(full, "rb") as f:
                raw = f.read().decode("ascii", errors="ignore")
            import re as _re
            texto = _re.sub(r"\\[a-z]+-?\d* ?|[{}]", "", raw)
        except Exception as e:
            return f"Error leyendo .rtf: {e}"
    else:
        return (f"Error: extensión '{ext}' no soportada por ingest_file. "
                f"Soportadas: .docx .pdf .txt .csv .log .tsv .json .yaml .rtf")

    if not texto or not texto.strip():
        return f"El archivo '{path}' no contiene texto extraíble (¿vacío o solo imágenes?)."

    if len(texto) > 50000:
        texto = texto[:50000] + "\n\n…[truncado en 50k chars]…"

    rel_orig = _osp.relpath(full, vroot).replace("\\", "/")
    if target_md:
        rel_md = target_md.lstrip("/")
        if not rel_md.endswith(".md"):
            rel_md += ".md"
    else:
        rel_md = rel_orig + ".md"

    iso = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
    md_body = (
        "---\n"
        f"source_file: {rel_orig}\n"
        f"source_type: {ext.lstrip('.')}\n"
        f"ingested_at: {iso}\n"
        "ingest_tool: ingest_file\n"
        "---\n\n"
        f"> Texto extraído automáticamente de `{rel_orig}` "
        f"({ext.lstrip('.').upper()}) para indexación semántica.\n\n"
        f"{texto}\n"
    )

    rel_md_url = rel_md.replace(" ", "%20")
    url = f"{OBSIDIAN_URL}/vault/{rel_md_url}"
    try:
        r = requests.put(url, headers=HEADERS, data=md_body.encode("utf-8"), timeout=15)
        if r.status_code in (200, 201, 204):
            return (f"OK: ingerido '{rel_orig}' → '{rel_md}' ({len(texto)} chars). "
                    f"Smart Connections lo indexará en segundos.")
        log.warning(f"⚠️ ingest_file({path}) → HTTP {r.status_code}: {r.text[:200]}")
        return f"Fallo al crear .md gemelo. HTTP {r.status_code}: {r.text[:120]}"
    except Exception as e:
        return f"Error de conexión: {e}"


def find_similar_notes(query: str, k: int = 5) -> str:
    """Búsqueda semántica en el vault vía el plugin mcp-tools (Jack Steam).

    mcp-tools registra la ruta POST /search/smart SOBRE el mismo Obsidian Local
    REST API (puerto 27123) que ya usamos para leer/escribir notas, así que
    reutilizamos OBSIDIAN_URL y el Bearer token. NO necesita un puerto aparte.

    Requiere: plugin mcp-tools + smart-connections habilitados en el vault.
    Devuelve los top-k pasajes (ruta + score + extracto truncado).
    """
    url = f"{OBSIDIAN_URL}/search/smart"
    payload = {"query": query, "filter": {"limit": int(k)}}
    # IMPORTANTE (verificado 29/05): mcp-tools recibe el body como TEXTO crudo y lo
    # hace JSON.parse internamente (arktype). Si enviamos Content-Type application/json,
    # el Obsidian Local REST API ya lo parsea a objeto y la validación falla con
    # HTTP 400 "must be a string (was an object)". Por eso mandamos el JSON
    # serializado como texto (text/plain) vía `data`, NO con `json=`.
    headers = {
        "Authorization": HEADERS["Authorization"],
        "Content-Type": "text/plain",
        "Accept": "application/json",
    }
    try:
        r = requests.post(
            url,
            headers=headers,
            data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
            timeout=20,
        )
    except Exception as e:
        return f"Error de conexión con búsqueda semántica: {str(e)}"

    if r.status_code == 503:
        return ("Error: el plugin mcp-tools o smart-connections no está activo en "
                "Obsidian. Habilítalos y reinicia Obsidian.")
    if r.status_code == 404:
        return ("Error: la ruta /search/smart no existe. ¿Está habilitado el plugin "
                "mcp-tools en este vault?")
    if r.status_code != 200:
        return f"Error en búsqueda semántica: HTTP {r.status_code}"

    try:
        results = r.json().get("results", [])
    except Exception as e:
        return f"Error parseando resultados de búsqueda: {str(e)}"

    if not results:
        return f"Sin resultados semánticos para: '{query}'."

    # Truncamos el texto de cada nota para no inflar el contexto del LLM.
    compact = []
    for item in results[: int(k)]:
        texto = (item.get("text") or "").strip()
        if len(texto) > 800:
            texto = texto[:800] + "…"
        compact.append({
            "path": item.get("path"),
            "score": round(item.get("score", 0), 4),
            "breadcrumbs": item.get("breadcrumbs"),
            "extracto": texto,
        })
    return json.dumps(compact, ensure_ascii=False, indent=2)


def run_template(name: str, arguments: Optional[dict] = None, target_path: str = "") -> str:
    """Ejecuta una plantilla de Templater vía mcp-tools (/templates/execute).

    A diferencia de /search/smart, este endpoint SÍ espera Content-Type
    application/json (verificado 29/05). En la plantilla, los argumentos se leen
    con `<% tp.mcpTools.prompt("clave") %>`.

    - name: ruta de la plantilla en el vault (ej '99_TEMPLATES/ficha_personaje.md').
    - arguments: dict {clave: valor} que rellena la plantilla.
    - target_path: si se indica, crea la nota ahí con el resultado; si se omite,
      solo devuelve el texto renderizado (sin escribir nada).

    Requiere los plugins Templater + mcp-tools habilitados.
    """
    url = f"{OBSIDIAN_URL}/templates/execute"
    payload = {"name": name, "arguments": arguments or {}}
    if target_path:
        if not target_path.endswith(".md"):
            target_path += ".md"
        payload["createFile"] = True
        payload["targetPath"] = target_path
    headers = {
        "Authorization": HEADERS["Authorization"],
        "Content-Type": "application/json",
        "Accept": "application/json",
    }
    try:
        r = requests.post(url, headers=headers, json=payload, timeout=30)
    except Exception as e:
        return f"Error de conexión con plantillas: {str(e)}"

    if r.status_code == 503:
        return ("Error: el plugin Templater no está activo en Obsidian. Habilítalo y "
                "reinicia Obsidian.")
    if r.status_code == 404:
        return f"Error: plantilla '{name}' no encontrada. Revisa la ruta."
    if r.status_code != 200:
        return f"Error ejecutando plantilla: HTTP {r.status_code} {r.text[:200]}"

    try:
        data = r.json()
    except Exception:
        return "Error parseando la respuesta de la plantilla."

    content = data.get("content", "")
    if target_path:
        return f"¡Éxito! Plantilla ejecutada y nota creada en {target_path}.\n\n{content}"
    return content or data.get("message", "Plantilla ejecutada sin contenido.")


def _git(args: list, timeout: int = 30):
    """Ejecuta `git -C <vault> <args>` de forma segura (sin shell). Devuelve (rc, salida)."""
    try:
        p = subprocess.run(
            [GIT_EXE, "-C", VAULT_PATH] + args,
            capture_output=True, text=True, timeout=timeout,
        )
        return p.returncode, ((p.stdout or "") + (p.stderr or "")).strip()
    except FileNotFoundError:
        return 127, "git no está instalado o no está en el PATH del sistema."
    except Exception as e:
        return 1, f"Error ejecutando git: {e}"


def git_status() -> str:
    """Muestra el estado git del vault: rama, cambios pendientes y último commit."""
    rc, _ = _git(["rev-parse", "--is-inside-work-tree"])
    if rc != 0:
        return "El vault no es un repositorio git todavía."
    _, branch = _git(["rev-parse", "--abbrev-ref", "HEAD"])
    _, changes = _git(["status", "--short"])
    _, last = _git(["log", "-1", "--pretty=format:%h %s (%cr)"])
    pendientes = changes if changes else "(sin cambios pendientes)"
    return f"Rama: {branch}\nÚltimo commit: {last}\nCambios pendientes:\n{pendientes}"


def git_save(message: str = "") -> str:
    """Guarda una versión del vault: git add -A + commit con el mensaje dado.
    Si no se da mensaje, usa uno por defecto. Si no hay cambios, lo indica."""
    rc, _ = _git(["rev-parse", "--is-inside-work-tree"])
    if rc != 0:
        return "El vault no es un repositorio git. No se puede guardar versión."
    if not message or not message.strip():
        message = "Copia de seguridad del libro"
    rc_add, out_add = _git(["add", "-A"])
    if rc_add != 0:
        return f"Error en git add: {out_add}"
    rc_c, out_c = _git([
        "-c", "user.name=AgenteEscritor", "-c", "user.email=agente@local",
        "commit", "-m", message,
    ])
    if rc_c != 0:
        if "nothing to commit" in out_c or "nada para hacer commit" in out_c:
            return "No hay cambios nuevos que guardar."
        return f"Error en commit: {out_c}"
    primera = out_c.splitlines()[0] if out_c else message
    return f"¡Versión guardada! {primera}"


def read_pdf(path: str, max_pages: int = 30) -> str:
    """Extrae el texto de un PDF del vault con pypdf (PDFs digitales).
    `path` es relativo a la raíz del vault o absoluto. Si el PDF está escaneado
    (sin texto extraíble), avisa de que necesitaría OCR."""
    import os.path as _osp
    full = _osp.realpath(path if _osp.isabs(path) else _osp.join(VAULT_PATH, path))
    vroot = _osp.realpath(VAULT_PATH)
    if not (full == vroot or full.startswith(vroot + os.sep)):
        return "Error: por seguridad solo puedo leer archivos dentro del vault."
    if not _osp.exists(full):
        return f"Error: no se encontró el PDF en '{path}'."
    try:
        from pypdf import PdfReader
    except Exception as e:
        return f"Error: la librería pypdf no está disponible ({e})."
    try:
        reader = PdfReader(full)
        n = len(reader.pages)
        partes = []
        for i, page in enumerate(reader.pages[:max_pages]):
            t = (page.extract_text() or "").strip()
            if t:
                partes.append(f"--- Página {i + 1} ---\n{t}")
        texto = "\n\n".join(partes)
        if not texto.strip():
            # PDF sin texto (escaneado) → fallback a Mistral OCR (tier gratuito ~1000 pág/mes)
            log.info(f"📄 PDF sin texto con pypdf, probando Mistral OCR: {path}")
            return _mistral_ocr_pdf(full, path)
        aviso = "" if n <= max_pages else f"\n\n[Truncado: {max_pages} de {n} páginas.]"
        if len(texto) > 12000:
            texto = texto[:12000] + "…"
        return f"PDF '{path}' ({n} páginas):\n\n{texto}{aviso}"
    except Exception as e:
        return f"Error leyendo el PDF: {e}"


def _mistral_ocr_pdf(full_path: str, display_path: str) -> str:
    """Fallback OCR para PDFs escaneados usando Mistral OCR (mistral-ocr-latest).
    Flujo verificado: files.upload(purpose=ocr) → get_signed_url → ocr.process."""
    if not MISTRAL_API_KEY:
        return "El PDF parece escaneado y no hay MISTRAL_API_KEY para hacer OCR."
    import os.path as _osp
    fname = _osp.basename(full_path)
    try:
        client = Mistral(api_key=MISTRAL_API_KEY)
        with open(full_path, "rb") as fh:
            up = client.files.upload(
                file={"file_name": fname, "content": fh}, purpose="ocr"
            )
        signed = client.files.get_signed_url(file_id=up.id)
        resp = client.ocr.process(
            document={"document_name": fname, "document_url": signed.url},
            model="mistral-ocr-latest",
            include_image_base64=False,
        )
        try:
            client.files.delete(file_id=up.id)
        except Exception:
            pass
        partes = []
        for page in resp.pages:
            md = (page.markdown or "").strip()
            if md:
                partes.append(f"--- Página {page.index + 1} ---\n{md}")
        texto = "\n\n".join(partes)
        if not texto.strip():
            return "El OCR no extrajo texto del PDF."
        if len(texto) > 12000:
            texto = texto[:12000] + "…"
        return f"PDF '{display_path}' (OCR Mistral, {len(resp.pages)} páginas):\n\n{texto}"
    except Exception as e:
        return f"Error en OCR Mistral: {e}"


def ocr_image(path: str) -> str:
    """Extrae el texto de una imagen (foto/escaneo) usando Mistral OCR.
    `path` relativo a la raíz del vault o absoluto. Soporta JPG, JPEG, PNG, WEBP, BMP.
    Ideal para fotos de apuntes a mano, capturas de pizarra o páginas escaneadas."""
    import os.path as _osp
    import base64 as _b64
    SUPPORTED = (".jpg", ".jpeg", ".png", ".webp", ".bmp")
    full = _osp.realpath(path if _osp.isabs(path) else _osp.join(VAULT_PATH, path))
    vroot = _osp.realpath(VAULT_PATH)
    if not (full == vroot or full.startswith(vroot + os.sep)):
        return "Error: por seguridad solo puedo leer archivos dentro del vault."
    if not _osp.exists(full):
        return f"Error: no se encontró el archivo en '{path}'."
    ext = _osp.splitext(full)[1].lower()
    if ext not in SUPPORTED:
        return f"Error: formato no soportado '{ext}'. Usa JPG, PNG, WEBP o BMP."
    if not MISTRAL_API_KEY:
        return "Error: falta MISTRAL_API_KEY en .env para hacer OCR de imágenes."
    try:
        mime = "image/jpeg" if ext in (".jpg", ".jpeg") else f"image/{ext.lstrip('.')}"
        with open(full, "rb") as fh:
            b64_data = _b64.b64encode(fh.read()).decode()
        client = Mistral(api_key=MISTRAL_API_KEY)
        resp = client.ocr.process(
            document={
                "type": "image_url",
                "image_url": f"data:{mime};base64,{b64_data}",
            },
            model="mistral-ocr-latest",
            include_image_base64=False,
        )
        partes = []
        for page in resp.pages:
            md = (page.markdown or "").strip()
            if md:
                partes.append(md)
        texto = "\n\n".join(partes)
        if not texto.strip():
            return "El OCR no extrajo texto de la imagen. Comprueba que la foto tenga buena iluminación y el texto sea legible."
        if len(texto) > 12000:
            texto = texto[:12000] + "…"
        return f"Texto extraído de '{path}' (OCR Mistral):\n\n{texto}"
    except Exception as e:
        return f"Error en OCR de imagen: {e}"


def _guardar_json(path: str, obj) -> None:
    """Guarda un dict como JSON sin romper si falla (best-effort)."""
    import json as _json
    try:
        with open(path, "w", encoding="utf-8") as fh:
            _json.dump(obj, fh, ensure_ascii=False)
    except Exception:
        pass


def _ocr_one(full_path: str) -> tuple:
    """OCR de UNA imagen (ruta absoluta ya validada).
    Devuelve (estado, texto_o_error). estado: 'ok' | 'rate' | 'error'."""
    import base64 as _b64, os.path as _osp
    ext = _osp.splitext(full_path)[1].lower()
    mime = "image/jpeg" if ext in (".jpg", ".jpeg") else f"image/{ext.lstrip('.')}"
    try:
        with open(full_path, "rb") as fh:
            b64_data = _b64.b64encode(fh.read()).decode()
        client = Mistral(api_key=MISTRAL_API_KEY)
        resp = client.ocr.process(
            document={"type": "image_url", "image_url": f"data:{mime};base64,{b64_data}"},
            model="mistral-ocr-latest",
            include_image_base64=False,
        )
        partes = [(p.markdown or "").strip() for p in resp.pages]
        return ("ok", "\n\n".join(x for x in partes if x))
    except Exception as e:
        msg = str(e)
        low = msg.lower()
        if any(k in low for k in ("429", "rate", "quota", "limit", "too many", "capacity")):
            return ("rate", msg)
        return ("error", msg)


def ocr_carpeta(folder: str, output_folder: str = "", pausa_seg: float = 2.0,
                recursivo: bool = True, crear_indice: bool = True) -> str:
    """Transcribe TODAS las imágenes de una carpeta del vault con OCR (lotes grandes).

    Pensado para cientos/miles de fotos de apuntes manuscritos.
    - Reanuda solo: si se corta, vuelve a llamarlo y salta lo ya hecho (.ocr_progreso.json).
    - Ritmo tranquilo: espera `pausa_seg` entre imágenes (por defecto 2s, sin prisa).
    - Si la API se satura (429/quota): espera 90s y reintenta; si sigue, PARA y avisa para reanudar.
    - Guarda cada imagen como .md en `output_folder` (por defecto subcarpeta `_transcripciones`).
    - Si crear_indice=True, crea una nota índice con enlaces [[...]] a todo lo transcrito.
    """
    import os.path as _osp, time as _time, json as _json
    if not MISTRAL_API_KEY:
        return "Error: falta MISTRAL_API_KEY en .env para hacer OCR."
    SUPPORTED = (".jpg", ".jpeg", ".png", ".webp", ".bmp")
    vroot = _osp.realpath(VAULT_PATH)
    src = _osp.realpath(folder if _osp.isabs(folder) else _osp.join(VAULT_PATH, folder))
    if not (src == vroot or src.startswith(vroot + os.sep)):
        return "Error: por seguridad solo puedo leer carpetas dentro del vault."
    if not _osp.isdir(src):
        return f"Error: la carpeta '{folder}' no existe."
    # Carpeta de salida (dentro del vault)
    if output_folder:
        out = _osp.realpath(output_folder if _osp.isabs(output_folder) else _osp.join(VAULT_PATH, output_folder))
    else:
        out = _osp.join(src, "_transcripciones")
    if not (out == vroot or out.startswith(vroot + os.sep)):
        return "Error: la carpeta de salida debe estar dentro del vault."
    try:
        os.makedirs(out, exist_ok=True)
    except Exception as e:
        return f"Error: no pude crear la carpeta de salida: {e}"

    # Listar imágenes (sin entrar en la carpeta de salida)
    imgs = []
    if recursivo:
        for root, _dirs, files in os.walk(src):
            if _osp.realpath(root).startswith(out):
                continue
            for f in files:
                if _osp.splitext(f)[1].lower() in SUPPORTED:
                    imgs.append(_osp.join(root, f))
    else:
        for f in os.listdir(src):
            full = _osp.join(src, f)
            if _osp.isfile(full) and _osp.splitext(f)[1].lower() in SUPPORTED:
                imgs.append(full)
    imgs.sort()
    if not imgs:
        return f"No encontré imágenes (JPG/PNG/WEBP/BMP) en '{folder}'."

    # Progreso (para reanudar)
    prog_path = _osp.join(out, ".ocr_progreso.json")
    done = {}
    if _osp.exists(prog_path):
        try:
            with open(prog_path, "r", encoding="utf-8") as fh:
                done = _json.load(fh)
        except Exception:
            done = {}

    total = len(imgs)
    nuevos, saltados, errores, transcritos = 0, 0, [], []
    for img in imgs:
        rel = _osp.relpath(img, vroot)
        if done.get(rel) == "ok":
            saltados += 1
            continue
        estado, texto = _ocr_one(img)
        if estado == "rate":
            _time.sleep(90)
            estado, texto = _ocr_one(img)
        if estado == "rate":
            _guardar_json(prog_path, done)
            pend = total - saltados - nuevos
            return (f"⏸️ La API de OCR está saturada (límite de uso temporal).\n"
                    f"Llevo {nuevos} transcritas nuevas y {saltados} ya estaban hechas. "
                    f"Quedan {pend} pendientes.\n"
                    f"Espera unos minutos y dime 'sigue con el OCR de {folder}' — "
                    f"reanudaré exactamente desde donde lo dejé.")
        if estado == "error":
            errores.append(f"{_osp.basename(img)}: {texto[:80]}")
            done[rel] = "error"
            continue
        # OK → guardar .md
        base = _osp.splitext(_osp.basename(img))[0]
        md_path = _osp.join(out, base.replace(" ", "_") + ".md")
        cuerpo = (texto or "").strip() or "*(El OCR no extrajo texto legible de esta imagen.)*"
        contenido = (
            f"---\norigen: \"{rel}\"\ntipo: transcripcion-ocr\n"
            f"fecha_ocr: {_time.strftime('%Y-%m-%d')}\nrevisado: false\n---\n\n"
            f"# {base}\n\n"
            f"> Transcripción automática (OCR) de `{_osp.basename(img)}`. "
            f"Puede tener errores de lectura — revísala.\n\n{cuerpo}\n"
        )
        try:
            with open(md_path, "w", encoding="utf-8") as fh:
                fh.write(contenido)
            done[rel] = "ok"
            nuevos += 1
            transcritos.append(base)
        except Exception as e:
            errores.append(f"{_osp.basename(img)}: no se pudo guardar ({e})")
            done[rel] = "error"
        if nuevos % 10 == 0:
            _guardar_json(prog_path, done)
        _time.sleep(max(0.0, pausa_seg))

    _guardar_json(prog_path, done)

    # Índice con wikilinks (semilla del grafo)
    indice_msg = ""
    if crear_indice and transcritos:
        lineas = [
            f"---\ntipo: indice-transcripciones\nactualizado: {_time.strftime('%Y-%m-%d')}\n---\n",
            f"# Índice de transcripciones — {_osp.basename(src)}\n",
            f"OCR automático, revisar. Total acumulado: {sum(1 for v in done.values() if v == 'ok')} imágenes.\n",
        ]
        lineas += [f"- [[{b}]]" for b in sorted(transcritos)]
        idx_path = _osp.join(out, "_INDICE_TRANSCRIPCIONES.md")
        try:
            with open(idx_path, "w", encoding="utf-8") as fh:
                fh.write("\n".join(lineas) + "\n")
            indice_msg = f"\n📑 Índice con enlaces: {_osp.relpath(idx_path, vroot)}"
        except Exception:
            pass

    resumen = (
        f"✅ OCR terminado en '{folder}'.\n"
        f"- Transcritas nuevas: {nuevos}\n"
        f"- Ya estaban hechas (saltadas): {saltados}\n"
        f"- Total de imágenes: {total}\n"
        f"- Guardadas en: {_osp.relpath(out, vroot)}/"
    )
    if errores:
        resumen += f"\n- ⚠️ {len(errores)} con error:\n  " + "\n  ".join(errores[:10])
        if len(errores) > 10:
            resumen += f"\n  …y {len(errores) - 10} más."
    resumen += indice_msg
    resumen += ("\n\nSiguiente paso sugerido: pídeme que cree fichas de wiki "
                "(personajes, lugares, hallazgos, eventos…) a partir de estas "
                "transcripciones para construir el grafo de conocimiento.")
    return resumen


def read_docx(path: str) -> str:
    """Extrae el texto de un .docx del vault. `path` relativo a la raíz del vault o absoluto
    (solo dentro del vault, por seguridad)."""
    import os.path as _osp
    full = _osp.realpath(path if _osp.isabs(path) else _osp.join(VAULT_PATH, path))
    vroot = _osp.realpath(VAULT_PATH)
    if not (full == vroot or full.startswith(vroot + os.sep)):
        return "Error: por seguridad solo puedo leer archivos dentro del vault."
    if not _osp.exists(full):
        return f"Error: no se encontró el documento en '{path}'."
    try:
        import docx
    except Exception as e:
        return f"Error: la librería python-docx no está disponible ({e})."
    try:
        d = docx.Document(full)
        partes = [p.text for p in d.paragraphs if p.text and p.text.strip()]
        texto = "\n".join(partes)
        if not texto.strip():
            return "El documento no contiene texto (¿vacío o solo imágenes?)."
        if len(texto) > 12000:
            texto = texto[:12000] + "…"
        return f"DOCX '{path}':\n\n{texto}"
    except Exception as e:
        return f"Error leyendo el .docx: {e}"


def fetch_url(url: str, max_chars: int = 8000) -> str:
    """Descarga una URL y devuelve su texto legible (sin HTML). Para leer un artículo o
    página concreta cuyo enlace ya se conoce (distinto de search_internet, que solo busca)."""
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    try:
        r = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=15)
        if r.status_code != 200:
            return f"Error: HTTP {r.status_code} al descargar {url}"
    except Exception as e:
        return f"Error de conexión: {e}"
    try:
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "noscript"]):
            tag.decompose()
        texto = " ".join(soup.get_text(separator=" ").split())
        if len(texto) > max_chars:
            texto = texto[:max_chars] + "…"
        return texto or "(la página no tiene texto legible)"
    except Exception as e:
        return f"Error procesando HTML: {e}"


# Plugins cuya config contiene API keys: NUNCA se leen ni editan vía estos tools.
_CONFIG_SECRET_PLUGINS = {"obsidian-local-rest-api", "bmo-chatbot"}


def _safe_plugin_id(plugin_id: str) -> bool:
    """True si el id es un nombre simple (sin traversal ni separadores)."""
    return bool(plugin_id) and not any(c in plugin_id for c in ("/", "\\", "..", "\x00"))


def list_obsidian_plugins() -> str:
    """Lista los plugins de Obsidian instalados con su versión y si están habilitados.
    Sirve para diagnosticar problemas de configuración."""
    import os.path as _osp
    base = _osp.join(VAULT_PATH, ".obsidian")
    plugins_dir = _osp.join(base, "plugins")
    try:
        cp = _osp.join(base, "community-plugins.json")
        enabled = json.loads(open(cp, encoding="utf-8").read()) if _osp.exists(cp) else []
    except Exception:
        enabled = []
    out = []
    try:
        for pid in sorted(os.listdir(plugins_dir)):
            mpath = _osp.join(plugins_dir, pid, "manifest.json")
            ver = "?"
            if _osp.exists(mpath):
                try:
                    ver = json.loads(open(mpath, encoding="utf-8").read()).get("version", "?")
                except Exception:
                    pass
            out.append({"id": pid, "version": ver, "enabled": pid in enabled})
    except Exception as e:
        return f"Error listando plugins: {e}"
    return json.dumps(out, ensure_ascii=False, indent=2)


def read_obsidian_config(plugin_id: str) -> str:
    """Lee la configuración (data.json) de un plugin de Obsidian para diagnosticar.
    Por seguridad NO devuelve configs con claves API (local-rest-api, bmo-chatbot)."""
    import os.path as _osp
    if not _safe_plugin_id(plugin_id):
        return "Error: id de plugin no válido."
    if plugin_id in _CONFIG_SECRET_PLUGINS:
        return (f"Por seguridad no se muestra la config de '{plugin_id}' (contiene claves API).")
    path = _osp.join(VAULT_PATH, ".obsidian", "plugins", plugin_id, "data.json")
    if not _osp.exists(path):
        return f"El plugin '{plugin_id}' no tiene data.json (o no existe)."
    try:
        content = open(path, encoding="utf-8").read()
    except Exception as e:
        return f"Error leyendo config: {e}"
    if len(content) > 8000:
        content = content[:8000] + "…"
    return f"Config de '{plugin_id}':\n{content}"


def update_obsidian_config(plugin_id: str, new_json: str) -> str:
    """Reescribe data.json de un plugin (CUIDADO). Valida que sea JSON, hace backup .bak y
    NUNCA toca plugins con claves API. Obsidian debe REINICIARSE para aplicar; si está
    abierto puede sobrescribir el cambio."""
    import os.path as _osp
    if not _safe_plugin_id(plugin_id) or plugin_id in _CONFIG_SECRET_PLUGINS:
        return f"Por seguridad no se puede editar la config de '{plugin_id}'."
    try:
        parsed = json.loads(new_json)  # validar
    except Exception as e:
        return f"Error: el contenido no es JSON válido ({e}). No se ha escrito nada."
    path = _osp.join(VAULT_PATH, ".obsidian", "plugins", plugin_id, "data.json")
    if not _osp.exists(path):
        return f"El plugin '{plugin_id}' no tiene data.json. No se crea uno nuevo por seguridad."
    try:
        with open(path, encoding="utf-8") as f:
            old = f.read()
        with open(path + ".bak", "w", encoding="utf-8") as f:
            f.write(old)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(parsed, f, ensure_ascii=False, indent=2)
        return (f"Config de '{plugin_id}' actualizada (backup en data.json.bak). "
                f"Reinicia Obsidian para aplicar.")
    except Exception as e:
        return f"Error escribiendo config: {e}"


_TEMPORAL_KEYWORDS = (
    "hoy", "actual", "actuales", "reciente", "recientes", "esta semana",
    "este mes", "este año", "noticias", "ultim", "últim", "latest",
    "today", "recent", "news", "yesterday", "ayer",
)


def _enrich_query_with_date(query: str) -> str:
    """Anclamos la query al año en curso para evitar que Mistral recupere
    noticias viejas de 2024. Reglas:
    - Si la query contiene años 2022-2024 explícitos, los sustituimos por año actual.
    - Si la query es temporal (palabras 'hoy', 'noticias', etc.) y no tiene año,
      añadimos el año actual al final.
    - Resto de queries, intactas."""
    current_year = datetime.now().year
    q = query

    # 1) Sustituir años 2022/2023/2024 por año actual si aparecen (Mistral los inventa)
    q = re.sub(r"\b(202[234])\b", str(current_year), q)

    # 2) Si no hay año y la query huele a temporal, añadir año actual
    has_year = bool(re.search(r"\b20\d{2}\b", q))
    lower = q.lower()
    is_temporal = any(kw in lower for kw in _TEMPORAL_KEYWORDS)
    if not has_year and is_temporal:
        q = f"{q} {current_year}"

    if q != query:
        log.info(f"🕐 Query reescrita para fecha actual: '{query}' -> '{q}'")
    return q


def _wrap_results_with_date_anchor(resultados_json: str) -> str:
    """Envuelve los resultados de búsqueda con un recordatorio de la fecha real
    para que Mistral NO reporte fechas que aparecen en los snippets antiguos."""
    fecha = _fecha_humana_es()
    año = datetime.now().year
    aviso = (
        f"[FECHA_REAL_DE_HOY: {fecha}. AÑO_EN_CURSO: {año}. "
        f"IMPORTANTE: los snippets pueden referirse a fechas antiguas, "
        f"pero cuando resumas al usuario indica SIEMPRE que es información "
        f"vigente a día de HOY ({año}). NO digas 'octubre 2024' ni 'el futuro', "
        f"estamos en {año}.]\n\nRESULTADOS:\n"
    )
    return aviso + resultados_json


def search_internet(query: str, max_results: int = 5) -> str:
    """Busca en internet con Tavily si hay key, si no DuckDuckGo HTML scraping.
    Reescribe la query para anclar al año actual y envuelve los resultados con
    un aviso de fecha real (anti-alucinación temporal de Mistral)."""
    query = _enrich_query_with_date(query)
    log.info(f"🔍 search_internet: '{query}'")

    # Camino A: Tavily (si hay key) — más fiable, con filtro temporal
    if TAVILY_API_KEY:
        try:
            r = requests.post(
                "https://api.tavily.com/search",
                json={
                    "api_key": TAVILY_API_KEY,
                    "query": query,
                    "max_results": max_results,
                    "search_depth": "basic",
                    "days": 30,  # solo resultados de los últimos 30 días
                },
                timeout=15,
            )
            if r.status_code == 200:
                data = r.json()
                payload = json.dumps(data.get("results", []), ensure_ascii=False, indent=2)
                return _wrap_results_with_date_anchor(payload)
            log.warning(f"Tavily HTTP {r.status_code}, fallback a DuckDuckGo")
        except Exception as e:
            log.warning(f"Tavily falló: {e}, fallback a DuckDuckGo")

    # Camino B: DuckDuckGo HTML scrape (sin API key, gratis)
    try:
        r = requests.get(
            "https://html.duckduckgo.com/html/",
            params={"q": query},
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
                    "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
                )
            },
            timeout=15,
        )
        if r.status_code != 200:
            return f"[search_internet] DuckDuckGo HTTP {r.status_code}"
        soup = BeautifulSoup(r.text, "html.parser")
        resultados = []
        for div in soup.select("div.result")[:max_results]:
            titulo_a = div.select_one("a.result__a")
            snippet_el = div.select_one(".result__snippet")
            url_el = div.select_one("a.result__url")
            if titulo_a:
                resultados.append({
                    "titulo": titulo_a.get_text(strip=True),
                    "url": (url_el.get_text(strip=True) if url_el else titulo_a.get("href", "")),
                    "snippet": snippet_el.get_text(strip=True) if snippet_el else "",
                })
        if not resultados:
            return f"[search_internet] Sin resultados para: '{query}'"
        payload = json.dumps(resultados, ensure_ascii=False, indent=2)
        return _wrap_results_with_date_anchor(payload)
    except Exception as e:
        log.error(f"search_internet error: {e}")
        return f"[search_internet] Error: {str(e)}"


# ==========================================
# 🎓 TOOLS DE INVESTIGACIÓN ACADÉMICA (las usa el perfil Investigador de Miguel)
# ==========================================

def buscar_papers(query: str, limite: int = 5) -> str:
    """Busca papers académicos REALES en OpenAlex (gratis, SIN API key). Devuelve título,
    autores, año y DOI verificables. Pensada para dar bibliografía sin inventar citas."""
    log.info(f"🎓 buscar_papers: '{query}'")
    try:
        lim = max(1, min(int(limite or 5), 10))
    except (TypeError, ValueError):
        lim = 5
    try:
        r = requests.get(
            "https://api.openalex.org/works",
            params={"search": query, "per-page": lim, "mailto": "agente@local"},
            timeout=25,
        )
        if r.status_code != 200:
            return f"[buscar_papers] OpenAlex HTTP {r.status_code}: {r.text[:160]}"
        results = r.json().get("results", [])
        if not results:
            return f"[buscar_papers] Sin resultados para: '{query}'"
        out = []
        for w in results[:lim]:
            title = w.get("title") or "(sin título)"
            year = w.get("publication_year") or "?"
            doi = w.get("doi") or ""
            url = doi or w.get("id") or ""
            authors = ", ".join(
                a.get("author", {}).get("display_name", "?")
                for a in (w.get("authorships") or [])[:4]
            ) or "(autores no listados)"
            cited = w.get("cited_by_count", 0)
            # OpenAlex da el abstract como índice invertido; reconstruimos un fragmento.
            abs_idx = w.get("abstract_inverted_index")
            abstract = ""
            if abs_idx:
                try:
                    pos = {}
                    for word, idxs in abs_idx.items():
                        for i in idxs:
                            pos[i] = word
                    abstract = " ".join(pos[i] for i in sorted(pos))[:280]
                except Exception:
                    abstract = ""
            entry = (f"• {title} ({year})\n  Autores: {authors}\n  DOI/URL: {url}\n  Citas: {cited}")
            if abstract:
                entry += f"\n  Resumen: {abstract}…"
            out.append(entry)
        return ("Resultados de OpenAlex (REALES y verificables — cita con el DOI, sin '→ VERIFICAR'):\n\n"
                + "\n\n".join(out))
    except Exception as e:
        log.error(f"buscar_papers error: {e}")
        return f"[buscar_papers] Error: {str(e)}"


def buscar_datasets(query: str, limite: int = 5) -> str:
    """Busca datasets, modelos 3D, preprints y publicaciones de acceso abierto en Zenodo
    (gratis, SIN API key). Útil para material primario digital y datos de excavación."""
    log.info(f"🎓 buscar_datasets: '{query}'")
    try:
        lim = max(1, min(int(limite or 5), 10))
    except (TypeError, ValueError):
        lim = 5
    try:
        r = requests.get(
            "https://zenodo.org/api/records",
            params={"q": query, "size": lim, "sort": "bestmatch"},
            timeout=25,
        )
        if r.status_code != 200:
            return f"[buscar_datasets] Zenodo HTTP {r.status_code}: {r.text[:160]}"
        hits = r.json().get("hits", {}).get("hits", [])
        if not hits:
            return f"[buscar_datasets] Sin resultados para: '{query}'"
        out = []
        for h in hits[:lim]:
            md = h.get("metadata", {})
            title = md.get("title", "(sin título)")
            year = str(md.get("publication_date", "?"))[:4]
            tipo = (md.get("resource_type", {}) or {}).get("type", "?")
            doi = h.get("doi") or md.get("doi") or ""
            url = h.get("links", {}).get("self_html") or (f"https://doi.org/{doi}" if doi else "")
            out.append(f"• {title} ({year}) · tipo: {tipo}\n  {url}")
        return ("Resultados de Zenodo (acceso abierto, verificables):\n\n" + "\n\n".join(out))
    except Exception as e:
        log.error(f"buscar_datasets error: {e}")
        return f"[buscar_datasets] Error: {str(e)}"


# ==========================================
# 📚 TOOLS DE ESTUDIO (resumen, esquema, quiz)
# ==========================================

def _tool_llm_call(system_prompt: str, user_prompt: str, max_tokens: int = 4096) -> str:
    """Llamada interna al LLM para tools que generan contenido.
    Usa provider/model del .env. SIN tools (use_tools=False) para evitar recursion."""
    provider = os.getenv("LLM_PROVIDER", "mistral").strip().lower()
    if provider not in PROVIDERS:
        provider = "mistral"
    model = PROVIDERS[provider]["default_model"]
    if provider == "mistral":
        model = MISTRAL_MODEL
    msgs = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    try:
        resp = _chat_completion(provider, model, msgs, use_tools=False)
        return resp["choices"][0]["message"].get("content", "")
    except Exception as e:
        log.error(f"_tool_llm_call error: {e}")
        return f"Error: el LLM no respondió ({e})"


def create_summary(nota_path: str, formato: str = "bullet") -> str:
    """Lee una nota del vault y genera un resumen via LLM.
    Formatos: 'bullet' (puntos), 'parrafo' (prosa), 'ficha' (ficha estudio con fechas+truco memoria).
    Guarda en 00_ESTUDIO/resumenes/Resumen_{nombre}.md"""
    content = read_obsidian_note(nota_path)
    if content.startswith("Error"):
        return f"[create_summary] No pude leer la nota: {content}"
    if len(content) > 8000:
        content = content[:8000] + "\n…(truncado)"

    formato = formato.strip().lower()
    if formato == "parrafo":
        fmt_instruction = "Escribe un resumen en prosa fluida (2-4 párrafos)."
    elif formato == "ficha":
        fmt_instruction = (
            "Escribe una ficha de estudio con: fechas clave, conceptos principales, "
            "un truco mnemotécnico para recordar lo esencial, y 3 preguntas de repaso."
        )
    else:
        formato = "bullet"
        fmt_instruction = "Escribe un resumen en puntos (bullet points) claros y concisos."

    system_prompt = (
        "Eres un asistente de estudio experto. Genera resúmenes precisos y útiles "
        "para preparar oposiciones. Usa markdown. Responde en español."
    )
    user_prompt = f"{fmt_instruction}\n\nContenido de la nota:\n\n{content}"

    resumen = _tool_llm_call(system_prompt, user_prompt)
    if resumen.startswith("Error:"):
        return f"[create_summary] {resumen}"

    nombre_base = Path(nota_path).stem
    now = datetime.now()
    md_body = (
        f"---\ntipo: resumen\nfuente: \"{nota_path}\"\nformato: {formato}\n"
        f"creado: {now.strftime('%Y-%m-%d %H:%M')}\n---\n\n"
        f"# Resumen: {nombre_base}\n\n{resumen}\n"
    )
    target = f"00_ESTUDIO/resumenes/Resumen_{nombre_base}.md"
    result = create_obsidian_note(target, md_body)
    if "Éxito" in result:
        return f"Resumen ({formato}) creado en `{target}`."
    return f"[create_summary] Resumen generado pero fallo al guardar: {result}"


def generate_esquema(tema: str, formato: str = "jerarquico") -> str:
    """Genera esquema jerarquico/timeline/comparativo basado en wiki del vault.
    NO usa mermaid, solo markdown con bullets indentados, negritas, [[wikilinks]].
    Guarda en 00_ESTUDIO/esquemas/Esquema_{tema_safe}.md"""
    context_raw = find_similar_notes(tema, k=5)
    context = context_raw[:4000] if len(context_raw) > 4000 else context_raw

    formato = formato.strip().lower()
    if formato == "timeline":
        fmt_instruction = (
            "Genera un esquema tipo línea temporal (timeline) con fechas y eventos "
            "ordenados cronológicamente. Usa bullets con fechas en negrita."
        )
    elif formato == "comparativo":
        fmt_instruction = (
            "Genera un esquema comparativo (tabla o columnas paralelas) que contraste "
            "los conceptos principales. Usa bullets indentados y negritas."
        )
    else:
        formato = "jerarquico"
        fmt_instruction = (
            "Genera un esquema jerárquico con niveles de indentación (bullets anidados). "
            "Destaca conceptos clave en **negrita** y enlaza notas con [[wikilinks]]."
        )

    system_prompt = (
        "Eres un asistente de estudio experto en crear esquemas visuales en markdown. "
        "NO uses mermaid ni diagramas de código. Usa solo bullets indentados, negritas "
        "y [[wikilinks]] de Obsidian. Responde en español."
    )
    user_prompt = (
        f"{fmt_instruction}\n\nTema: {tema}\n\n"
        f"Contexto del vault (notas relacionadas):\n{context}"
    )

    esquema = _tool_llm_call(system_prompt, user_prompt)
    if esquema.startswith("Error:"):
        return f"[generate_esquema] {esquema}"

    tema_safe = re.sub(r'[^\w\s-]', '', tema)[:60].strip().replace(' ', '_')
    now = datetime.now()
    md_body = (
        f"---\ntipo: esquema\ntema: \"{tema}\"\nformato: {formato}\n"
        f"creado: {now.strftime('%Y-%m-%d %H:%M')}\n---\n\n"
        f"# Esquema: {tema}\n\n{esquema}\n"
    )
    target = f"00_ESTUDIO/esquemas/Esquema_{tema_safe}.md"
    result = create_obsidian_note(target, md_body)
    if "Éxito" in result:
        return f"Esquema ({formato}) creado en `{target}`."
    return f"[generate_esquema] Esquema generado pero fallo al guardar: {result}"


def generate_quiz(tema: str, n_preguntas: int = 5, tipo: str = "mixto") -> str:
    """Genera cuestionario (test/desarrollo/mixto/verdadero_falso) con soluciones.
    Guarda en 00_ESTUDIO/ejercicios/Quiz_{tema}_{timestamp}.md"""
    context_raw = find_similar_notes(tema, k=5)
    context = context_raw[:6000] if len(context_raw) > 6000 else context_raw

    n_preguntas = max(1, min(int(n_preguntas), 20))
    tipo = tipo.strip().lower()
    if tipo == "test":
        tipo_instruction = (
            f"Genera {n_preguntas} preguntas tipo test (4 opciones, solo 1 correcta). "
            "Al final incluye las soluciones con explicación breve."
        )
    elif tipo == "desarrollo":
        tipo_instruction = (
            f"Genera {n_preguntas} preguntas de desarrollo (respuesta larga). "
            "Al final incluye una respuesta modelo para cada una."
        )
    elif tipo == "verdadero_falso":
        tipo_instruction = (
            f"Genera {n_preguntas} afirmaciones de verdadero/falso. "
            "Al final incluye las soluciones indicando cuáles son verdaderas y cuáles falsas, con justificación."
        )
    else:
        tipo = "mixto"
        tipo_instruction = (
            f"Genera {n_preguntas} preguntas mezclando tipos: test (4 opciones), "
            "desarrollo corto y verdadero/falso. Al final incluye todas las soluciones."
        )

    system_prompt = (
        "Eres un preparador de oposiciones experto. Genera cuestionarios rigurosos "
        "y pedagógicos basados en el material proporcionado. Incluye siempre las "
        "soluciones al final, separadas de las preguntas. Usa markdown. Responde en español."
    )
    user_prompt = (
        f"{tipo_instruction}\n\nTema: {tema}\n\n"
        f"Material de referencia del vault:\n{context}"
    )

    quiz = _tool_llm_call(system_prompt, user_prompt)
    if quiz.startswith("Error:"):
        return f"[generate_quiz] {quiz}"

    tema_safe = re.sub(r'[^\w\s-]', '', tema)[:60].strip().replace(' ', '_')
    now = datetime.now()
    timestamp = now.strftime('%Y%m%d_%H%M')
    md_body = (
        f"---\ntipo: quiz\ntema: \"{tema}\"\nformato: {tipo}\n"
        f"n_preguntas: {n_preguntas}\ncreado: {now.strftime('%Y-%m-%d %H:%M')}\n---\n\n"
        f"# Quiz: {tema}\n\n{quiz}\n"
    )
    target = f"00_ESTUDIO/ejercicios/Quiz_{tema_safe}_{timestamp}.md"
    result = create_obsidian_note(target, md_body)
    if "Éxito" in result:
        return f"Quiz ({tipo}, {n_preguntas} preguntas) creado en `{target}`."
    return f"[generate_quiz] Quiz generado pero fallo al guardar: {result}"


# Mapeo para ejecución
names_to_functions = {
    "read_obsidian_note": read_obsidian_note,
    "create_obsidian_note": create_obsidian_note,
    "update_obsidian_note": update_obsidian_note,
    "overwrite_obsidian_note": overwrite_obsidian_note,
    "delete_obsidian_note": delete_obsidian_note,
    "list_vault_files": list_vault_files,
    "scan_disk_files": scan_disk_files,
    "ingest_file": ingest_file,
    "find_similar_notes": find_similar_notes,
    "run_template": run_template,
    "git_status": git_status,
    "git_save": git_save,
    "read_pdf": read_pdf,
    "ocr_image": ocr_image,
    "ocr_carpeta": ocr_carpeta,
    "read_docx": read_docx,
    "fetch_url": fetch_url,
    "list_obsidian_plugins": list_obsidian_plugins,
    "read_obsidian_config": read_obsidian_config,
    "update_obsidian_config": update_obsidian_config,
    "search_internet": search_internet,
    "create_summary": create_summary,
    "generate_esquema": generate_esquema,
    "generate_quiz": generate_quiz,
    "buscar_papers": buscar_papers,
    "buscar_datasets": buscar_datasets,
}

tools = [
     {
        "type": "function",
        "function": {
            "name": "read_obsidian_note",
            "description": "Lee el contenido exacto de una nota guardada en tu vault de Obsidian.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Nombre del archivo, ej: 'Personajes/Duque_Juan.md'"}
                },
                "required": ["filename"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_obsidian_note",
            "description": "Crea una nota física en el disco duro del vault de Obsidian. Úsalo si el usuario te pide generar una ficha, esquema o capitulo.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Ruta y nombre (Personajes/Nuevo.md)"},
                    "content": {"type": "string", "description": "El contenido markdown completo a escribir."}
                },
                "required": ["filename", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "update_obsidian_note",
            "description": "Añade texto al final de una nota existente.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Ruta y nombre"},
                    "content": {"type": "string", "description": "El texto que añadirás al final."}
                },
                "required": ["filename", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "overwrite_obsidian_note",
            "description": "Reescribe completamente una nota existente. CUIDADO: borra todo el contenido anterior. Úsalo solo si el usuario pide reescribir o reemplazar una nota entera.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Ruta y nombre de la nota a reescribir"},
                    "content": {"type": "string", "description": "El contenido markdown completo nuevo."}
                },
                "required": ["filename", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "delete_obsidian_note",
            "description": "Borra una nota del vault. La nota va a la papelera de Obsidian. Úsalo solo si el usuario pide explícitamente borrar una nota.",
            "parameters": {
                "type": "object",
                "properties": {
                    "filename": {"type": "string", "description": "Ruta y nombre de la nota a borrar"}
                },
                "required": ["filename"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_vault_files",
            "description": "Lista los archivos del vault de Obsidian (vía REST API). Por defecto lista TODAS las extensiones (md, docx, pdf, txt, etc.) para que veas archivos que el usuario haya podido pegar. OJO: respeta `showUnsupportedFiles` de Obsidian — si está en false, los docx/txt no aparecen aunque existan; entonces usa `scan_disk_files`.",
            "parameters": {
                "type": "object",
                "properties": {
                    "folder": {"type": "string", "description": "Carpeta a listar (ej: '03_NOTAS'). Por defecto '/' (raíz)."},
                    "extension": {"type": "string", "description": "Filtrar por extensión, vacío = todas. Admite coma: '.md,.docx,.pdf'."}
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "scan_disk_files",
            "description": "Lista archivos del vault leyendo DIRECTAMENTE del disco, sin pasar por la REST API de Obsidian. Útil cuando Obsidian oculta archivos no-soportados (.docx, .txt, etc.). Recorre recursivamente todas las subcarpetas.",
            "parameters": {
                "type": "object",
                "properties": {
                    "folder": {"type": "string", "description": "Subcarpeta del vault, vacío = raíz."},
                    "extension": {"type": "string", "description": "Filtro extensión (vacío = todas). Admite coma: '.md,.docx'."}
                },
                "required": [],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ingest_file",
            "description": "Extrae el texto de un archivo NO-md (.docx, .pdf, .txt, .csv, .rtf) del vault y crea una nota .md GEMELA al lado con el contenido. La nota .md la indexa Smart Connections automáticamente, así el usuario puede buscar semánticamente dentro de docs Word, PDFs, etc. ÚSALO siempre que detectes con `scan_disk_files` o `list_vault_files` un archivo no-md que aún NO tenga su gemelo .md.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Ruta del archivo original dentro del vault, ej '03_NOTAS/borrador.docx'."},
                    "target_md": {"type": "string", "description": "Ruta destino del .md (opcional). Si vacío, usa '{path}.md'."}
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "find_similar_notes",
            "description": "Busca pasajes del vault semánticamente similares a una consulta (búsqueda por significado, no por nombre exacto). Úsalo cuando el usuario pregunte por un tema, personaje, lugar o idea sin saber en qué archivo está, o para encontrar capítulos/notas relacionadas.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "Consulta en lenguaje natural (tema, personaje, idea)"},
                    "k": {"type": "integer", "description": "Número de resultados a devolver (por defecto 5)"}
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "run_template",
            "description": "Ejecuta una plantilla de Templater para generar una nota estructurada (ficha de personaje, lugar, etc.) y opcionalmente la crea en el vault. Plantillas disponibles: '99_TEMPLATES/ficha_personaje.md' (claves: nombre, descripcion, rol, aspecto, relaciones, notas) y '99_TEMPLATES/ficha_lugar.md' (claves: nombre, descripcion, importancia, atmosfera, notas). Úsalo cuando el usuario pida crear una ficha de personaje o lugar.",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "Ruta de la plantilla, ej '99_TEMPLATES/ficha_personaje.md'"},
                    "arguments": {"type": "object", "description": "Diccionario clave:valor que rellena la plantilla (las claves dependen de la plantilla)", "additionalProperties": {"type": "string"}},
                    "target_path": {"type": "string", "description": "Ruta donde crear la nota resultante, ej '02_WIKI/personajes/Иван.md'. Si se omite, solo devuelve el texto sin crear archivo."}
                },
                "required": ["name", "arguments"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "git_status",
            "description": "Muestra el estado del control de versiones (git) del vault: rama actual, cambios sin guardar y último commit. Úsalo cuando el usuario pregunte si hay cambios sin guardar o por el historial.",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "git_save",
            "description": "Guarda una versión (commit) de TODO el vault con un mensaje descriptivo. Úsalo cuando el usuario pida 'guarda', 'haz copia de seguridad', 'versiona' o tras un trabajo importante. Escribe un mensaje claro de qué cambió.",
            "parameters": {
                "type": "object",
                "properties": {
                    "message": {"type": "string", "description": "Mensaje del commit describiendo qué se guardó, ej 'Añadido capítulo 3 y ficha de Iván'"}
                },
                "required": ["message"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_pdf",
            "description": "Extrae el texto de un archivo PDF del vault para leer material de referencia. Úsalo cuando el usuario mencione un PDF o pida resumir/consultar un documento PDF.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Ruta del PDF dentro del vault, ej '05_REFERENCIAS/historia.pdf'"},
                    "max_pages": {"type": "integer", "description": "Máximo de páginas a leer (por defecto 30)"}
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ocr_image",
            "description": "Extrae el texto de una imagen (foto o escaneo) usando OCR. Úsalo cuando el usuario suba una foto de apuntes, una captura de pizarra, o cualquier imagen con texto. Soporta JPG, PNG, WEBP, BMP. Si el usuario dice 'lee mis apuntes' y hay imágenes en 01_CRUDO/, úsalo.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Ruta de la imagen dentro del vault, ej '01_CRUDO/apuntes_tema5.jpg'"}
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "ocr_carpeta",
            "description": "Transcribe con OCR TODAS las imágenes de una carpeta (lotes grandes: cientos o miles de fotos de apuntes manuscritos). Reanuda solo si se corta, va con pausas para no saturar la API, y si la API se satura PARA y avisa para reanudar luego. Guarda cada imagen como una nota .md y crea un índice con enlaces. Úsalo cuando el usuario diga 'procesa/transcribe todas mis notas/apuntes/fotos' de una carpeta.",
            "parameters": {
                "type": "object",
                "properties": {
                    "folder": {"type": "string", "description": "Carpeta del vault con las imágenes, ej '01_CRUDO/apuntes_mano'"},
                    "output_folder": {"type": "string", "description": "Carpeta donde guardar los .md (opcional; por defecto una subcarpeta '_transcripciones')"},
                    "pausa_seg": {"type": "number", "description": "Segundos de espera entre imágenes (opcional, por defecto 2). Súbelo si la API se satura."}
                },
                "required": ["folder"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_docx",
            "description": "Extrae el texto de un archivo Word (.docx) del vault. Úsalo cuando el usuario mencione un documento Word o un borrador en .docx.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Ruta del .docx dentro del vault, ej '01_CRUDO/borrador.docx'"}
                },
                "required": ["path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_url",
            "description": "Descarga y devuelve el texto legible de una página web concreta cuyo enlace ya conoces. Distinto de search_internet (que busca). Úsalo para leer un artículo, una entrada de Wikipedia, etc.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "La URL completa a leer"}
                },
                "required": ["url"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_obsidian_plugins",
            "description": "Lista los plugins de Obsidian instalados, su versión y si están habilitados. Úsalo para diagnosticar problemas de configuración del vault.",
            "parameters": {"type": "object", "properties": {}, "required": []},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_obsidian_config",
            "description": "Lee la configuración (data.json) de un plugin de Obsidian para diagnosticar un problema. No muestra configs con claves API por seguridad.",
            "parameters": {
                "type": "object",
                "properties": {
                    "plugin_id": {"type": "string", "description": "Id del plugin, ej 'templater-obsidian' o 'smart-connections'"}
                },
                "required": ["plugin_id"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "update_obsidian_config",
            "description": "Reescribe la configuración (data.json) de un plugin de Obsidian. CUIDADO: cambia ajustes del vault. Valida el JSON y hace backup. Úsalo SOLO si el usuario pide arreglar una config concreta. Avisa de que hay que reiniciar Obsidian.",
            "parameters": {
                "type": "object",
                "properties": {
                    "plugin_id": {"type": "string", "description": "Id del plugin a configurar"},
                    "new_json": {"type": "string", "description": "El contenido JSON completo y válido de la nueva config"}
                },
                "required": ["plugin_id", "new_json"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "search_internet",
            "description": "Usa internet para consultar información histórica o externa que no sabes. No lo uses si solo es corrección de textos locales.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "La consulta de investigación precisa"}
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "buscar_papers",
            "description": "Busca artículos académicos REALES en OpenAlex (gratis). Devuelve título, autores, año y DOI verificables. Úsalo para encontrar y citar bibliografía científica/histórica sin inventar referencias. Ideal para arqueología, historia y humanidades.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "Tema o palabras clave en inglés, ej 'Sea Peoples Late Bronze Age collapse'"},
                    "limite": {"type": "integer", "description": "Número de resultados (1-10, por defecto 5)"}
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "buscar_datasets",
            "description": "Busca datasets, modelos 3D de artefactos, preprints y publicaciones de acceso abierto en Zenodo (gratis). Útil para datos de excavación y material primario digital.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "Tema o palabras clave, ej 'Hittite excavation pottery'"},
                    "limite": {"type": "integer", "description": "Número de resultados (1-10, por defecto 5)"}
                },
                "required": ["query"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "create_summary",
            "description": "Lee una nota del vault y genera un resumen automático (bullet points, prosa o ficha de estudio). Lo guarda en 00_ESTUDIO/resumenes/.",
            "parameters": {
                "type": "object",
                "properties": {
                    "nota_path": {"type": "string", "description": "Ruta de la nota a resumir, ej: 'Capitulo_3.md'"},
                    "formato": {"type": "string", "enum": ["bullet", "parrafo", "ficha"], "description": "Formato del resumen: bullet (puntos), parrafo (prosa), ficha (estudio con trucos memoria). Por defecto: bullet."}
                },
                "required": ["nota_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_esquema",
            "description": "Genera un esquema de estudio (jerárquico, timeline o comparativo) sobre un tema, usando notas del vault como contexto. Lo guarda en 00_ESTUDIO/esquemas/.",
            "parameters": {
                "type": "object",
                "properties": {
                    "tema": {"type": "string", "description": "El tema sobre el que generar el esquema"},
                    "formato": {"type": "string", "enum": ["jerarquico", "timeline", "comparativo"], "description": "Tipo de esquema: jerarquico (bullets anidados), timeline (cronológico), comparativo (contraste). Por defecto: jerarquico."}
                },
                "required": ["tema"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "generate_quiz",
            "description": "Genera un cuestionario de estudio con preguntas y soluciones sobre un tema, usando notas del vault como referencia. Lo guarda en 00_ESTUDIO/ejercicios/.",
            "parameters": {
                "type": "object",
                "properties": {
                    "tema": {"type": "string", "description": "El tema sobre el que generar el cuestionario"},
                    "n_preguntas": {"type": "integer", "description": "Número de preguntas (1-20). Por defecto: 5."},
                    "tipo": {"type": "string", "enum": ["test", "desarrollo", "mixto", "verdadero_falso"], "description": "Tipo de preguntas: test (4 opciones), desarrollo, mixto, verdadero_falso. Por defecto: mixto."}
                },
                "required": ["tema"],
            },
        },
    }
]


# ==========================================
# 🔀 ENDPOINT PROXY (Engañando a BMO / Copilot)
# ==========================================

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

_MESES_ES = [
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]
_DIAS_ES = ["lunes", "martes", "miércoles", "jueves", "viernes", "sábado", "domingo"]


def _fecha_humana_es() -> str:
    now = datetime.now()
    return (
        f"{_DIAS_ES[now.weekday()]}, {now.day} de {_MESES_ES[now.month - 1]} "
        f"de {now.year}, {now.strftime('%H:%M')}"
    )


def build_critical_context() -> str:
    """Contexto mínimo que siempre se inyecta: recordatorio de tools disponibles.
    No inyectamos fecha: si el modelo la necesita, que llame a `search_internet`."""
    return (
        "[HERRAMIENTAS DISPONIBLES PARA ESTE CHAT]\n"
        "- `search_internet(query)`: búsqueda web en tiempo real. Úsala SIEMPRE que "
        "necesites información actual (fecha de hoy, noticias, eventos, cotizaciones, "
        "clima, datos posteriores a tu entrenamiento). No inventes, busca.\n"
        "- `create_obsidian_note(filename, content)`: crea una nota .md nueva en el vault.\n"
        "- `read_obsidian_note(filename)`: lee el contenido de una nota existente.\n"
        "- `update_obsidian_note(filename, content)`: añade texto al final de una nota.\n"
        "- `overwrite_obsidian_note(filename, content)`: reescribe una nota entera (DESTRUCTIVO).\n"
        "- `delete_obsidian_note(filename)`: borra una nota (va a papelera Obsidian, DESTRUCTIVO).\n"
        "- `list_vault_files(folder, extension)`: lista archivos del vault (REST API). Por "
        "defecto lista TODAS las extensiones. Si esperabas ver un .docx/.txt y no aparece, "
        "Obsidian lo oculta — usa `scan_disk_files` que lee disco directo.\n"
        "- `scan_disk_files(folder, extension)`: lista archivos leyendo el disco (no pasa "
        "por Obsidian). Útil para detectar archivos pegados manualmente que Obsidian no muestra.\n"
        "- `ingest_file(path)`: convierte un archivo NO-md (.docx, .pdf, .txt, .rtf) en un .md "
        "gemelo. ÚSALO siempre que veas un archivo no-md sin gemelo .md para que Smart "
        "Connections pueda indexarlo y permitir búsqueda semántica sobre su contenido.\n"
        "- `find_similar_notes(query, k)`: búsqueda SEMÁNTICA del vault. Úsala cuando el "
        "usuario pregunte por un tema/personaje/lugar/idea sin saber el nombre exacto del "
        "archivo, o para encontrar pasajes relacionados. Prefiérela a list_vault_files "
        "cuando la búsqueda sea por significado, no por nombre.\n"
        "- `run_template(name, arguments, target_path)`: genera una nota estructurada con "
        "una plantilla Templater. Úsalo para fichas de personaje "
        "('99_TEMPLATES/ficha_personaje.md') o lugar ('99_TEMPLATES/ficha_lugar.md') en "
        "vez de escribir el markdown a mano.\n"
        "- `git_status()`: muestra cambios sin guardar e historial del vault.\n"
        "- `git_save(message)`: guarda una versión (commit) del vault. Úsalo cuando el "
        "usuario pida guardar o hacer copia de seguridad.\n"
        "- `read_pdf(path)`: extrae el texto de un PDF del vault (con OCR automático si está escaneado).\n"
        "- `ocr_image(path)`: extrae el texto de una imagen (foto/escaneo) con OCR. Úsalo cuando "
        "el usuario suba fotos de apuntes a 01_CRUDO/ (JPG, PNG, WEBP, BMP).\n"
        "- `ocr_carpeta(folder, pausa_seg)`: transcribe con OCR TODAS las imágenes de una carpeta "
        "(lotes grandes, cientos/miles de fotos). Reanuda solo si se corta y avisa si la API se "
        "satura. Úsalo cuando el usuario diga 'transcribe/procesa todas mis notas o fotos' de una carpeta.\n"
        "- `read_docx(path)`: extrae el texto de un documento Word (.docx) del vault.\n"
        "- `fetch_url(url)`: lee el texto de una página web concreta (cuyo enlace ya conoces).\n"
        "- `list_obsidian_plugins()` / `read_obsidian_config(plugin_id)`: diagnostican la "
        "configuración de Obsidian. `update_obsidian_config(plugin_id, new_json)` la modifica "
        "(solo si el usuario lo pide; avisa de reiniciar Obsidian).\n"
        "- `create_summary(nota_path, formato)`: genera un resumen automático de una nota "
        "(formatos: bullet, parrafo, ficha). Lo guarda en 00_ESTUDIO/resumenes/.\n"
        "- `generate_esquema(tema, formato)`: genera un esquema de estudio sobre un tema "
        "(formatos: jerarquico, timeline, comparativo). Usa notas del vault como contexto. "
        "Lo guarda en 00_ESTUDIO/esquemas/.\n"
        "- `generate_quiz(tema, n_preguntas, tipo)`: genera un cuestionario con soluciones "
        "(tipos: test, desarrollo, mixto, verdadero_falso). Usa notas del vault como "
        "referencia. Lo guarda en 00_ESTUDIO/ejercicios/.\n"
        "\n[REGLAS DE USO DE HERRAMIENTAS — OBLIGATORIAS]\n"
        "- Para crear, leer, editar o guardar notas DEBES LLAMAR a la herramienta correspondiente "
        "(create_obsidian_note, read_obsidian_note, update_obsidian_note, etc.). No describas lo que "
        "harías: HAZLO llamando a la herramienta.\n"
        "- create_obsidian_note REQUIERE 2 argumentos: filename (ruta, ej '03_NOTAS/mi_ficha.md') y "
        "content (el markdown). read_obsidian_note REQUIERE filename. Nunca llames a una herramienta "
        "sin sus argumentos obligatorios.\n"
        "- La autenticación de la REST API de Obsidian YA está configurada en el sistema. NUNCA pidas "
        "al usuario una API key, ni sugieras desactivar plugins, ni hables de 'Local REST API'.\n"
        "- PROHIBIDO inventar errores: NUNCA digas que una herramienta falló o dio un error (401, "
        "autorización, etc.) si NO la has llamado realmente en esta respuesta. Solo reporta un error "
        "si la herramienta lo devolvió de verdad, y entonces cítalo literal.\n"
        "Responde en el idioma del usuario (español o búlgaro).\n"
    )


SYSTEM_PROMPT_FALLBACK = (
    "Eres un agente de investigación y redacción con acceso a internet y al vault "
    "Obsidian del usuario. Sé conciso, útil y preciso. Si no sabes algo actual, "
    "busca con `search_internet` antes de responder."
)

@app.get("/v1/models")
async def get_models():
    """OpenAI-compatible. Anuncia 'agente-escritor' (default) + un modelo por proveedor,
    para que el SELECTOR DE MODELOS de BMO permita cambiar sin tocar .env. El usuario elige
    'proveedor:modelo' y el proxy enruta (cada proveedor necesita su key en .env)."""
    ids = ["agente-escritor", "yayo"]
    for prov, cfg in PROVIDERS.items():
        if cfg.get("local"):
            continue  # ollama/llamacpp: SIGUEN funcionando si se escriben, pero NO se anuncian en el menú
        ids.append(f"{prov}:{cfg['default_model']}")
    ids += [
        # Mistral extras
        "mistral:mistral-small-latest",
        # OpenAI extras
        "openai:gpt-4o",
        # DeepSeek extras
        "deepseek:deepseek-reasoner",
        # OpenRouter premium — IDs verificados 2026-06-06
        "openrouter:anthropic/claude-sonnet-4.6",
        "openrouter:anthropic/claude-haiku-4.5",
        "openrouter:openai/gpt-4o",
        "openrouter:google/gemini-2.5-flash",
        # Gemini extras (IDs reales verificados 2026-06-06)
        "gemini:gemini-2.5-flash",
        "gemini:gemini-3-flash-preview",
        "gemini:gemini-3.1-flash-lite",
        # OpenRouter GRATIS (con tools) — IDs verificados funcionando 2026-06-06
        "openrouter:google/gemma-4-31b-it:free",
        "openrouter:openai/gpt-oss-120b:free",
        "openrouter:meta-llama/llama-3.3-70b-instruct:free",
        # Groq (gratis, con tool calling) — verificados funcionando 2026-06-06
        "groq:llama-3.3-70b-versatile",
        "groq:openai/gpt-oss-120b",
        "groq:openai/gpt-oss-20b",
        "groq:qwen/qwen3-32b",
    ]
    # Dedup conservando orden (evita duplicados en el menú)
    seen, uniq = set(), []
    for i in ids:
        if i not in seen:
            seen.add(i)
            uniq.append(i)
    data = [
        {"id": i, "object": "model", "created": 1686935002, "owned_by": "agente-escritor"}
        for i in uniq
    ]
    return {"object": "list", "data": data}


@app.get("/health")
async def health():
    """Health check del proxy y de los servicios externos."""
    obsidian_ok = False
    obsidian_err = None
    try:
        r = requests.get(
            f"{OBSIDIAN_URL}/",
            headers={"Authorization": HEADERS["Authorization"]},
            timeout=3,
        )
        obsidian_ok = r.status_code == 200
        if not obsidian_ok:
            obsidian_err = f"HTTP {r.status_code}"
    except Exception as e:
        obsidian_err = str(e)

    return {
        "proxy": "ok",
        "port": PROXY_PORT,
        "mistral_key_present": bool(MISTRAL_API_KEY),
        "obsidian_url": OBSIDIAN_URL,
        "obsidian_reachable": obsidian_ok,
        "obsidian_error": obsidian_err,
        "tavily_present": bool(TAVILY_API_KEY),
        "model": MISTRAL_MODEL,
    }


@app.get("/diagnose")
async def diagnose():
    """Diagnóstico E2E: prueba GET, PUT, POST, DELETE contra la Obsidian REST API.
    Acceder con: curl http://localhost:9000/diagnose"""
    results = {}
    auth_hdr = {"Authorization": f"Bearer {OBSIDIAN_API_KEY}"}
    test_file = "__diag_test_deleteme.md"
    test_url = f"{OBSIDIAN_URL}/vault/{test_file}"
    test_body = "# Diagnóstico\nEsta nota la creó el proxy para verificar la REST API."

    # 1) GET raíz (sin auth) — debe dar 200 (es exempt)
    try:
        r = requests.get(f"{OBSIDIAN_URL}/", timeout=3)
        results["1_GET_root_no_auth"] = {"status": r.status_code, "ok": r.status_code == 200}
    except Exception as e:
        results["1_GET_root_no_auth"] = {"error": str(e)}

    # 2) GET vault con solo Authorization — debe funcionar (como read_obsidian_note)
    try:
        r = requests.get(f"{OBSIDIAN_URL}/vault/", headers=auth_hdr, timeout=3)
        results["2_GET_vault_auth_only"] = {"status": r.status_code, "ok": r.status_code == 200}
    except Exception as e:
        results["2_GET_vault_auth_only"] = {"error": str(e)}

    # 3) PUT con HEADERS compartido (como create_obsidian_note — EL QUE FALLA)
    try:
        r = requests.put(test_url, headers=HEADERS, data=test_body.encode("utf-8"), timeout=5)
        results["3_PUT_shared_HEADERS"] = {
            "status": r.status_code,
            "ok": r.status_code in [200, 201, 204],
            "body": r.text[:200],
            "sent_headers": dict(HEADERS),
        }
    except Exception as e:
        results["3_PUT_shared_HEADERS"] = {"error": str(e)}

    # 4) PUT con headers construidos en el momento (¿funciona diferente?)
    try:
        fresh_headers = {
            "Authorization": f"Bearer {OBSIDIAN_API_KEY}",
            "Content-Type": "text/markdown",
        }
        r = requests.put(test_url, headers=fresh_headers, data=test_body.encode("utf-8"), timeout=5)
        results["4_PUT_fresh_headers"] = {
            "status": r.status_code,
            "ok": r.status_code in [200, 201, 204],
            "body": r.text[:200],
        }
    except Exception as e:
        results["4_PUT_fresh_headers"] = {"error": str(e)}

    # 5) PUT SIN Content-Type (¿es ese el problema?)
    try:
        r = requests.put(test_url, headers=auth_hdr, data=test_body.encode("utf-8"), timeout=5)
        results["5_PUT_no_content_type"] = {
            "status": r.status_code,
            "ok": r.status_code in [200, 201, 204],
            "body": r.text[:200],
        }
    except Exception as e:
        results["5_PUT_no_content_type"] = {"error": str(e)}

    # 6) POST (como update_obsidian_note / append)
    try:
        r = requests.post(test_url, headers=HEADERS, data="\nAppend test.".encode("utf-8"), timeout=5)
        results["6_POST_append"] = {
            "status": r.status_code,
            "ok": r.status_code in [200, 201, 204],
            "body": r.text[:200],
        }
    except Exception as e:
        results["6_POST_append"] = {"error": str(e)}

    # 7) Limpieza: DELETE
    try:
        r = requests.delete(test_url, headers=auth_hdr, timeout=3)
        results["7_DELETE_cleanup"] = {"status": r.status_code}
    except Exception as e:
        results["7_DELETE_cleanup"] = {"error": str(e)}

    # 8) Dump de config para diagnóstico
    results["config"] = {
        "OBSIDIAN_URL": OBSIDIAN_URL,
        "OBSIDIAN_API_KEY_first8": OBSIDIAN_API_KEY[:8] + "..." if OBSIDIAN_API_KEY else "EMPTY",
        "HEADERS_keys": list(HEADERS.keys()),
    }

    return results

# ==========================================
# 🔀 MULTI-MODELO — proveedores OpenAI-compatibles
# ==========================================
# Todos exponen POST {base_url}/chat/completions con el MISMO formato (incl. tools y
# tool_calls), así que un solo cliente con `requests` sirve para todos. Mistral es el
# único validado con el bucle de tools; los demás se activan poniendo su key en .env.
PROVIDERS = {
    # --- Nube (necesitan API key en .env) ---
    "mistral":    {"base_url": "https://api.mistral.ai/v1", "key_env": "MISTRAL_API_KEY", "default_model": "mistral-large-latest"},
    "groq":       {"base_url": "https://api.groq.com/openai/v1", "key_env": "GROQ_API_KEY", "default_model": "openai/gpt-oss-120b"},
    "gemini":     {"base_url": "https://generativelanguage.googleapis.com/v1beta/openai", "key_env": "GEMINI_API_KEY", "default_model": "gemini-flash-latest"},
    "grok":       {"base_url": "https://api.x.ai/v1", "key_env": "XAI_API_KEY", "default_model": "grok-2-latest"},
    "openai":     {"base_url": "https://api.openai.com/v1", "key_env": "OPENAI_API_KEY", "default_model": "gpt-4o-mini"},
    "deepseek":   {"base_url": "https://api.deepseek.com/v1", "key_env": "DEEPSEEK_API_KEY", "default_model": "deepseek-chat"},
    "openrouter": {"base_url": "https://openrouter.ai/api/v1", "key_env": "OPENROUTER_API_KEY", "default_model": "openai/gpt-4o-mini"},
    # Claude vía endpoint OpenAI-compatible de Anthropic (beta). Las tools pueden diferir.
    "claude":     {"base_url": "https://api.anthropic.com/v1", "key_env": "ANTHROPIC_API_KEY", "default_model": "claude-sonnet-4-6"},
    # Kimi (Moonshot). Usa .ai para internacional, .cn para China.
    "kimi":       {"base_url": "https://api.moonshot.ai/v1", "key_env": "MOONSHOT_API_KEY", "default_model": "kimi-k2-0905-preview"},
    # MiniMax: endpoint OpenAI-compat; algunos planes exigen GroupId — verificar si falla.
    "minimax":    {"base_url": "https://minimax-m2.com/api/v1", "key_env": "MINIMAX_API_KEY", "default_model": "MiniMax-M2.7"},
    # --- Local (sin key; el usuario avanzado levanta el servidor) ---
    # Ollama: `ollama serve` expone OpenAI-compat en 11434. llama.cpp: usar LLAMACPP_BASE_URL.
    "ollama":     {"base_url": os.getenv("OLLAMA_BASE_URL", "http://localhost:11434/v1"), "key_env": "OLLAMA_API_KEY", "default_model": os.getenv("OLLAMA_MODEL", "mistral-local:latest"), "local": True},
    "llamacpp":   {"base_url": os.getenv("LLAMACPP_BASE_URL", "http://localhost:8081/v1"), "key_env": "LLAMACPP_API_KEY", "default_model": os.getenv("LLAMACPP_MODEL", "local-model"), "local": True},
}
DEFAULT_PROVIDER = (os.getenv("LLM_PROVIDER", "mistral").strip().lower() or "mistral")


def resolve_provider_and_model(requested_model: str):
    """Decide proveedor+modelo. Acepta 'proveedor:modelo', 'proveedor/modelo', un nombre
    de proveedor suelto, o cae al default (env LLM_PROVIDER, o mistral)."""
    rm = (requested_model or "").strip()
    for sep in (":", "/"):
        if sep in rm:
            prov, _, mod = rm.partition(sep)
            prov = prov.strip().lower()
            if prov in PROVIDERS:
                return prov, (mod.strip() or PROVIDERS[prov]["default_model"])
    if rm.lower() in PROVIDERS:
        return rm.lower(), PROVIDERS[rm.lower()]["default_model"]
    prov = DEFAULT_PROVIDER if DEFAULT_PROVIDER in PROVIDERS else "mistral"
    if prov == "mistral":
        return prov, MISTRAL_MODEL
    return prov, PROVIDERS[prov]["default_model"]


def _sanitize_messages(msgs: list) -> list:
    """Limpia el array de mensajes antes de enviarlo al proveedor:
    1. Fusiona mensajes assistant consecutivos (BMO a veces genera esto al filtrar historial)
    2. Asegura que el último mensaje no-system sea user/tool (Mistral lo exige)
    3. Elimina mensajes con content vacío"""
    if not msgs:
        return msgs
    cleaned = []
    for m in msgs:
        role = m.get("role", "")
        content = m.get("content")
        # Saltar mensajes vacíos (excepto si tienen tool_calls)
        if not content and not m.get("tool_calls") and role != "tool":
            continue
        # Fusionar assistant consecutivos
        if role == "assistant" and cleaned and cleaned[-1].get("role") == "assistant" and not cleaned[-1].get("tool_calls"):
            prev_content = cleaned[-1].get("content", "") or ""
            new_content = content or ""
            cleaned[-1] = {**cleaned[-1], "content": (prev_content + "\n" + new_content).strip()}
        else:
            cleaned.append(m)
    # Si el último mensaje es assistant (sin tool_calls), añadir user sintético
    if cleaned and cleaned[-1].get("role") == "assistant" and not cleaned[-1].get("tool_calls"):
        cleaned.append({"role": "user", "content": "Continúa."})
    return cleaned


def _chat_completion(provider: str, model: str, msgs: list, use_tools: bool = True, max_tokens: int = 4096) -> dict:
    """Llama a {base_url}/chat/completions (formato OpenAI) vía requests. Devuelve el JSON.
    Lanza RuntimeError si falta la key o si la respuesta no es 200.

    `max_tokens` se reenvía SIEMPRE (por defecto 4096): si no se manda, OpenRouter (y otros)
    reservan el máximo del modelo (p.ej. 16384) y rechazan con 402 si la key tiene poco saldo."""
    cfg = PROVIDERS[provider]
    api_key = os.getenv(cfg["key_env"], "").strip()
    if not api_key and not cfg.get("local"):
        # La key puede haberse añadido/cambiado en el .env DESPUÉS de arrancar el .exe.
        # Recargamos el .env en caliente (override) y reintentamos antes de fallar.
        # Así añadir una key funciona sin tener que reiniciar el .bat.
        try:
            for _cand in (".env", ".env.backend"):
                _p = APP_DIR / _cand
                if _p.exists():
                    load_dotenv(_p, override=True)
                    break
        except Exception:
            pass
        api_key = os.getenv(cfg["key_env"], "").strip()
    if not api_key:
        if cfg.get("local"):
            api_key = "local"  # Ollama/llama.cpp no requieren key real
        else:
            raise RuntimeError(f"Falta la API key {cfg['key_env']} para el proveedor '{provider}'.")
    url = cfg["base_url"].rstrip("/") + "/chat/completions"
    payload = {"model": model, "messages": msgs}
    try:
        if max_tokens and int(max_tokens) > 0:
            payload["max_tokens"] = int(max_tokens)
    except (TypeError, ValueError):
        payload["max_tokens"] = 4096
    if use_tools:
        payload["tools"] = tools
        payload["tool_choice"] = "auto"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    # Reintento ante 429 (rate-limit). Mistral free = ~1 req/seg; el bucle de agente puede
    # dispararse rápido y chocar. Esperamos y reintentamos en vez de romper la tarea.
    import time as _t
    r = None
    for _intento in range(3):
        r = requests.post(url, headers=headers, json=payload, timeout=120)
        if r.status_code == 429 and _intento < 2:
            log.warning(f"[429 {provider}] rate-limit, espero y reintento ({_intento + 1}/2)…")
            _t.sleep(3 * (_intento + 1))   # backoff: 3s, luego 6s
            continue
        break
    if r.status_code != 200:
        raise RuntimeError(f"HTTP {r.status_code} de {provider}: {r.text[:300]}")
    return r.json()


@app.post("/v1/chat/completions")
async def proxy_chat(request: Request):
    """Endpoint OpenAI-compatible. Acepta messages, opcionalmente tools (las ignora,
    nosotros forzamos las nuestras), retorna SIEMPRE JSON válido aunque haya error
    para evitar el famoso `SyntaxError: Unexpected token 'B', 'Bearer tok...'` en BMO."""
    try:
        data = await request.json()
    except Exception as e:
        log.error(f"Body no es JSON válido: {e}")
        return JSONResponse(
            status_code=400,
            content={"error": {"message": f"Invalid JSON body: {e}", "type": "invalid_request"}},
        )

    incoming_messages = data.get("messages", [])
    if not incoming_messages:
        return JSONResponse(
            status_code=400,
            content={"error": {"message": "messages array missing", "type": "invalid_request"}},
        )

    # (La validación de API key se hace por proveedor tras resolverlo, más abajo.)

    # SIEMPRE inyectamos el contexto crítico (fecha real + recordatorio de tools).
    # Si el cliente (Copilot/BMO) ya mandó un system, lo combinamos: contexto crítico
    # primero, luego su system. Así respetamos sus instrucciones pero garantizamos que
    # Mistral conozca la fecha y sepa que puede usar tools.
    critical = build_critical_context()
    if incoming_messages and incoming_messages[0].get("role") == "system":
        client_system = incoming_messages[0].get("content", "") or ""
        merged_system = (
            critical
            + "\n\n[INSTRUCCIONES DEL CLIENTE]\n"
            + client_system
        )
        messages = [{"role": "system", "content": merged_system}] + list(incoming_messages[1:])
    else:
        messages = [
            {"role": "system", "content": critical + "\n\n" + SYSTEM_PROMPT_FALLBACK}
        ] + list(incoming_messages)

    provider, model = resolve_provider_and_model(data.get("model", ""))
    messages = _sanitize_messages(messages)
    # max_tokens del cliente (BMO manda el del perfil, ~4096). Si no viene, 4096.
    # Esto evita que OpenRouter reserve el máximo del modelo (16384) y dé 402 con poco saldo.
    client_max_tokens = data.get("max_tokens") or 4096
    log.info(f"🚀 Petición interceptada ({len(messages)} mensajes). Proveedor={provider} · modelo={model} · max_tokens={client_max_tokens}")

    try:
        resp = _chat_completion(provider, model, messages, max_tokens=client_max_tokens)
    except Exception as e:
        emsg = str(e)
        low = emsg.lower()
        # 1) Límite por MINUTO del proveedor (típico de Groq gratis: 8K TPM). NO es contexto.
        if any(k in low for k in ("tokens per minute", "tpm", "rate_limit", "rate limit",
                                  "requests per minute", "rpm", "too many requests", "429")):
            log.warning(f"[rate-limit · {provider} · {model}] {emsg[:200]}")
            aviso = (
                f"⏳ El proveedor **{provider}** está saturado (límite de uso por minuto). "
                "No es que el texto sea largo: el plan gratis de Groq solo admite ~8.000 tokens/minuto "
                "y el agente con sus herramientas ya ocupa casi eso.\n\n"
                "Solución: usa **gemini:gemini-flash-latest** o **mistral:mistral-large-latest** "
                "(sin ese tope), o espera un minuto y reintenta."
            )
            return JSONResponse(
                status_code=200,
                content={"choices": [{"index": 0, "finish_reason": "stop",
                                      "message": {"role": "assistant", "content": aviso}}],
                         "model": model},
            )
        # 2) Texto demasiado largo para el CONTEXTO del modelo → mensaje claro, no "fallo inesperado".
        if any(k in low for k in ("context length", "context_length", "maximum context",
                                  "too many tokens", "tokens_limit", "input is too long",
                                  "string too long", "3230", "request entity too large",
                                  "reduce the length", "context window")):
            log.warning(f"[contexto excedido · {provider} · {model}] {emsg[:200]}")
            aviso = (
                f"⚠️ El texto es demasiado largo para el modelo actual (**{model}**). "
                "No he perdido nada de la conversación. Tienes 2 opciones:\n\n"
                "1. Pídemelo **por partes** (por ejemplo, un capítulo cada vez), o\n"
                "2. Cambia a un modelo de **contexto grande** (Gemini ~1M, o mistral-large) "
                "en el selector de modelos de BMO y repite."
            )
            return JSONResponse(
                status_code=200,
                content={
                    "choices": [{"index": 0, "finish_reason": "stop",
                                 "message": {"role": "assistant", "content": aviso}}],
                    "model": model,
                },
            )
        log.error(f"[LLM inicial · {provider}] {emsg}")
        return JSONResponse(
            status_code=502,
            content={"error": {"message": f"Error del proveedor {provider}: {emsg}", "type": "upstream_error"}},
        )

    msg = resp["choices"][0]["message"]
    messages.append(msg)

    # --- BUCLE DE AGENTE AUTÓNOMO (formato OpenAI uniforme para todos los proveedores) ---
    max_steps = 6
    step = 0
    while msg.get("tool_calls") and step < max_steps:
        step += 1
        tcs = msg["tool_calls"]
        log.info(f"[AGENTE paso {step}] ejecutando {len(tcs)} tool(s)…")

        for tc in tcs:
            fn = tc.get("function", {})
            f_name = fn.get("name", "")
            try:
                f_args = json.loads(fn.get("arguments") or "{}")
            except Exception:
                f_args = {}
            try:
                _args_repr = json.dumps(f_args, ensure_ascii=False)
                if len(_args_repr) > 300:
                    _args_repr = _args_repr[:300] + "…"
            except Exception:
                _args_repr = str(f_args)[:300]
            log.info(f"  └─ {f_name}({_args_repr})")
            f_func = names_to_functions.get(f_name)
            try:
                f_res = f_func(**f_args) if f_func else f"Función '{f_name}' no mapeada."
            except Exception as e:
                f_res = f"Error ejecutando {f_name}: {e}"
                log.error(f_res)
            # Log del resultado del tool (truncado a 200 chars para no saturar)
            _res_repr = str(f_res)[:200]
            if "Error" in _res_repr or "Fallo" in _res_repr or "error" in _res_repr:
                log.warning(f"  ⚠️ {f_name} → {_res_repr}")
            else:
                log.info(f"  ✔️ {f_name} → {_res_repr}")

            messages.append({
                "role": "tool",
                "name": f_name,
                "content": str(f_res),
                "tool_call_id": tc.get("id"),
            })

        try:
            resp = _chat_completion(provider, model, messages, max_tokens=client_max_tokens)
        except Exception as e:
            log.error(f"[LLM paso {step} · {provider}] {e}")
            return JSONResponse(
                status_code=502,
                content={"error": {"message": f"Error del proveedor {provider} en bucle: {e}", "type": "upstream_error"}},
            )
        msg = resp["choices"][0]["message"]
        messages.append(msg)

    final_text = msg.get("content") or "Tarea finalizada exitosamente."
    log.info("✅ Respuesta final generada.")

    return {
        "id": f"chatcmpl-{int(time.time())}",
        "object": "chat.completion",
        "created": int(time.time()),
        "model": data.get("model", "agente-escritor"),
        "choices": [
            {
                "index": 0,
                "message": {"role": "assistant", "content": final_text},
                "finish_reason": "stop",
            }
        ],
        "usage": {
            "prompt_tokens": 0,
            "completion_tokens": 0,
            "total_tokens": 0,
        },
    }


if __name__ == "__main__":
    import uvicorn
    log.info(f"🚀 Arrancando uvicorn en {PROXY_HOST}:{PROXY_PORT} ...")
    try:
        uvicorn.run(app, host=PROXY_HOST, port=PROXY_PORT)
    except Exception as e:
        log.error(f"❌ Uvicorn falló al arrancar: {e}")
        import traceback
        log.error(traceback.format_exc())
        # Mantener la ventana abierta para que el usuario vea el error
        input("Pulsa ENTER para cerrar...")
